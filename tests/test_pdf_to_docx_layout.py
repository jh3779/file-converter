"""PDF→DOCX 절대 위치 레이아웃 테스트 — DEC-037.

핵심 검증: 줄 단위로 원본과 같은 절대 위치(w:framePr, 페이지 좌표 고정)로
재구성되는지 — pdf_to_pptx(DEC-030)의 위치 재구성과 같은 원리를 DOCX의
레거시 프레임 기능으로 구현한 것. 손으로 만든 최소 PDF(hand-rolled content
stream)를 쓰는 관례는 test_pdf_to_pptx.py의 _mini_pdf와 동일 — 별도 PDF
생성 라이브러리 불요.
"""
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app import converters
from app.converters.base import ConversionError

EMU_PER_PT = 12700
TWIPS_PER_PT = 20


def _mini_pdf(path: Path, pages: list[list[tuple]], page_sizes: list[tuple] | None = None):
    """손으로 만든 최소 PDF. pages: 페이지별 [(text, x, y, font_key, size)] 목록.
    font_key "F1"=Helvetica(일반), "F2"=Helvetica-Bold(굵게 감지용) —
    test_pdf_to_pptx.py의 동명 헬퍼와 동일한 구조. page_sizes: 페이지별
    (width, height) — 생략 시 전부 612x792(레터)."""
    page_obj_start = 3
    n = len(pages)
    sizes = page_sizes or [(612, 792)] * n
    font1_num = page_obj_start + 2 * n
    font2_num = font1_num + 1
    kids, page_bodies, content_bodies = [], [], []
    for i, lines in enumerate(pages):
        page_idx = page_obj_start + 2 * i
        content_idx = page_idx + 1
        kids.append(f"{page_idx} 0 R")
        ops = []
        for text, x, y, font_key, size in lines:
            ops.append(f"BT /{font_key} {size} Tf {x} {y} Td ({text}) Tj ET")
        content = "\n".join(ops).encode()
        stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
        w, h = sizes[i]
        page_bodies.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {w} {h}] /Contents {content_idx} 0 R"
            f" /Resources << /Font << /F1 {font1_num} 0 R /F2 {font2_num} 0 R >> >> >>".encode())
        content_bodies.append(stream)
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        ("<< /Type /Pages /Kids [" + " ".join(kids) + f"] /Count {n} >>").encode(),
    ]
    for pb, cb in zip(page_bodies, content_bodies):
        objs.append(pb)
        objs.append(cb)
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


def _frame_pr(paragraph):
    """문단의 w:framePr 속성을 {x,y,w,h}(pt 단위)로 돌려준다. 없으면 None."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    fp = pPr.find(qn("w:framePr"))
    if fp is None:
        return None
    return {k: int(fp.get(qn(f"w:{k}"))) / TWIPS_PER_PT for k in ("x", "y", "w", "h")}


class TestPdfToDocxLayout(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_text_content_preserved(self):
        src = self.tmp / "single.pdf"
        _mini_pdf(src, [[("Hello world", 50, 700, "F1", 12)]])
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("Hello world", texts)

    def test_line_positioned_via_frame_pr(self):
        """줄 bbox(PDF 포인트, 원점 좌하단) → framePr(원점 좌상단) 변환이
        맞는지 — x는 그대로, y는 페이지높이-y1 근방이어야 한다."""
        src = self.tmp / "pos.pdf"
        _mini_pdf(src, [[("Positioned", 100, 700, "F1", 12)]])
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        p = next(p for p in doc.paragraphs if p.text == "Positioned")
        frame = _frame_pr(p)
        self.assertIsNotNone(frame, "framePr가 없음 — 절대 위치 지정이 안 됨")
        self.assertAlmostEqual(frame["x"], 100, delta=2)
        expected_y_approx = 792 - 700
        self.assertAlmostEqual(frame["y"], expected_y_approx, delta=15)

    def test_bold_font_detected_via_fontname_heuristic(self):
        """pdf_to_pptx(DEC-030)와 같은 휴리스틱 재사용 — Helvetica-Bold로
        그려진 줄은 run.font.bold=True로 반영돼야 한다."""
        src = self.tmp / "bold.pdf"
        _mini_pdf(src, [[
            ("Normal", 50, 700, "F1", 12),
            ("Bold", 50, 680, "F2", 12),
        ]])
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        bold_flags = {}
        for p in doc.paragraphs:
            for run in p.runs:
                bold_flags[run.text] = run.font.bold
        self.assertEqual(bold_flags.get("Normal"), False)
        self.assertEqual(bold_flags.get("Bold"), True)

    def test_page_size_matches_pdf_page_size(self):
        src = self.tmp / "size.pdf"
        _mini_pdf(src, [[("x", 50, 700, "F1", 12)]])
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        section = doc.sections[0]
        self.assertEqual(section.page_width, round(612 * EMU_PER_PT))
        self.assertEqual(section.page_height, round(792 * EMU_PER_PT))

    def test_multi_page_uses_page_breaks(self):
        src = self.tmp / "multi.pdf"
        _mini_pdf(src, [
            [("Page one", 50, 700, "F1", 12)],
            [("Page two", 50, 700, "F1", 12)],
        ])
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("Page one", texts)
        self.assertIn("Page two", texts)
        # 두 페이지를 나누는 명시적 페이지 나눔이 있어야 한다(w:br type="page").
        has_page_break = any(
            br.get(qn("w:type")) == "page"
            for p in doc.paragraphs
            for br in p._p.findall(f".//{qn('w:br')}")
        )
        self.assertTrue(has_page_break)

    def test_mixed_page_sizes_each_section_matches_its_page(self):
        """페이지마다 크기가 다른 PDF(스캔 첨부문서의 가로/세로 혼합 등)에서
        각 페이지가 자기 실제 크기의 섹션에 들어가야 한다 — 첫 페이지
        크기로 섹션이 고정된 채면 vAnchor="page" 기준 y좌표가 실제 렌더링
        페이지 높이와 어긋나 텍스트가 밀려 보인다(PR 콘텐츠 리뷰로 발견해
        수정)."""
        src = self.tmp / "mixed.pdf"
        _mini_pdf(
            src,
            [[("Portrait", 50, 700, "F1", 12)], [("Landscape", 50, 500, "F1", 12)]],
            page_sizes=[(612, 792), (792, 612)],
        )
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        self.assertEqual(len(doc.sections), 2)
        s0, s1 = doc.sections
        self.assertEqual(s0.page_width, round(612 * EMU_PER_PT))
        self.assertEqual(s0.page_height, round(792 * EMU_PER_PT))
        self.assertEqual(s1.page_width, round(792 * EMU_PER_PT))
        self.assertEqual(s1.page_height, round(612 * EMU_PER_PT))

        p2 = next(p for p in doc.paragraphs if p.text == "Landscape")
        frame = _frame_pr(p2)
        self.assertIsNotNone(frame)
        # 두 번째 페이지 실제 높이(612)를 기준으로 계산돼야 한다 —
        # 첫 페이지 높이(792)를 썼다면 180pt(=792-612)만큼 어긋난다.
        expected_y_approx = 612 - 500
        self.assertAlmostEqual(frame["y"], expected_y_approx, delta=15)

    def test_corrupted_pdf_rejected(self):
        src = self.tmp / "broken.pdf"
        src.write_bytes(b"not a real pdf")
        with self.assertRaises(ConversionError) as ctx:
            converters.convert(src, "docx", self.tmp)
        self.assertEqual(ctx.exception.key, "err.corrupted")


if __name__ == "__main__":
    unittest.main()
