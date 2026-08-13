"""PDF→DOCX 밑줄 감지 테스트 — 밑줄 감지 개선.

PDF는 밑줄을 폰트 속성이 아니라 별도의 벡터 선(그림)으로 그리는 경우가
많아(DEC-027), 굵게/기울임과 같은 폰트 이름 휴리스틱으로는 판별할 수 없다.
텍스트 바로 아래의 거의 수평인 벡터 선(LTLine)을 밑줄 후보로 보고, 글자
bbox와의 위치·겹침 관계로 실제 밑줄인지 판별한다(app/converters/pdf.py의
_underline_candidates/_char_is_underlined). 표준 14 폰트(Helvetica)는
/Widths 없이도 pdfminer 내장 AFM 메트릭으로 실제 글자 너비가 나오므로
(로컬 검증 확인 — 폰트 리소스 오브젝트 번호만 정확히 참조하면 됨), 이 파일의
손조립 PDF도 test_pdf_to_docx_visuals.py와 같은 방식을 그대로 쓴다.
"""
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app import converters

_HELVETICA = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
_FONT_RESOURCES = "/Font << /F1 5 0 R >>"


def _build_pdf(path: Path, content_ops: str, width=612, height=792):
    """단일 페이지 최소 PDF, 폰트 리소스(Helvetica)를 오브젝트 5로 고정
    배치 — test_pdf_to_docx_visuals.py의 _build_pdf와 동일한 조립 방식."""
    content = content_ops.encode()
    stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {width} {height}] /Contents 4 0 R"
         f" /Resources << {_FONT_RESOURCES} >> >>").encode(),
        stream,
        _HELVETICA,
    ]
    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


class TestPdfToDocxUnderline(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_word_with_line_beneath_marked_underline(self):
        """"underlined plain"에서 앞 단어 아래에만 짧은 수평선을 그으면
        그 단어의 run만 underline=True, 뒤 단어는 False여야 한다."""
        src = self.tmp / "underline.pdf"
        _build_pdf(src,
                   "BT /F1 12 Tf 100 600 Td (underlined plain) Tj ET "
                   "0 0 0 RG 1 w 100 598 m 156 598 l S")
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        by_text = {}
        for p in doc.paragraphs:
            for r in p.runs:
                by_text[r.text] = bool(r.font.underline)
        self.assertTrue(by_text.get("underlined"))
        self.assertFalse(by_text.get(" plain", True))

    def test_no_nearby_line_means_no_underline(self):
        """벡터 선이 전혀 없으면 밑줄이 감지되지 않아야 한다(오탐 방지 회귀)."""
        src = self.tmp / "plain.pdf"
        _build_pdf(src, "BT /F1 12 Tf 100 600 Td (plain text) Tj ET")
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        for p in doc.paragraphs:
            for r in p.runs:
                self.assertFalse(bool(r.font.underline))

    def test_distant_line_not_mistaken_for_underline(self):
        """텍스트에서 멀리 떨어진 선(표 테두리 등)은 밑줄로 오판되지 않고
        여전히 표 테두리 도형(w:pBdr)으로 렌더링돼야 한다."""
        src = self.tmp / "distant.pdf"
        _build_pdf(src,
                   "BT /F1 12 Tf 100 600 Td (no underline here) Tj ET "
                   "0 0 1 RG 1 w 100 500 m 300 500 l S")
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        for p in doc.paragraphs:
            for r in p.runs:
                self.assertFalse(bool(r.font.underline))
        bordered = [p for p in doc.paragraphs
                    if p._p.find(qn("w:pPr")) is not None
                    and p._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None]
        self.assertEqual(len(bordered), 1)

    def test_underline_consumed_not_double_rendered_as_border(self):
        """밑줄로 쓰인 선은 run.underline=True로 이미 반영됐으니 표 테두리
        도형으로 다시 그려 겹쳐 보이면 안 된다 — 밑줄로 소비된 선은
        w:pBdr 문단으로 남지 않아야 한다."""
        src = self.tmp / "consumed.pdf"
        _build_pdf(src,
                   "BT /F1 12 Tf 100 600 Td (underlined) Tj ET "
                   "0 0 0 RG 1 w 100 598 m 156 598 l S")
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        bordered = [p for p in doc.paragraphs
                    if p._p.find(qn("w:pPr")) is not None
                    and p._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None]
        self.assertEqual(len(bordered), 0)


if __name__ == "__main__":
    unittest.main()
