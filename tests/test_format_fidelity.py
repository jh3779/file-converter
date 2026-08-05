"""HWP→DOCX·PDF→DOCX 문자 서식(굵게/기울임/크기/색상) 반영 테스트 — DEC-027.

PDF 쪽은 hwplib/JDK 없이도 항상 실행되도록 손으로 만든 최소 PDF(폰트
리소스 이름에 Bold/Italic이 들어간 것)를 쓴다(LibreOffice 불필요).
HWP 쪽은 실제 hwplib 라이브러리가 필요해 기존 TestHwp와 동일한 조건에서만
실행한다 — 로컬 spike 빌드가 없는 CI "test" job에서는 스킵된다(HWP 서식
검증의 실제 CI 게이트는 scripts/smoke_pdf_pipeline.py와 build.yml의 HWP
엔진 스모크 쪽에 별도로 있음).
"""
import shutil
import subprocess
import tempfile
import unittest
from pathlib import Path

from docx import Document

from app import converters

REPO = Path(__file__).resolve().parents[1]
HWP_DISTRIBUTION = REPO / "spike" / "hwplib" / "repo" / "sample_hwp" / "distribution.hwp"


def _mini_pdf_with_bold(path: Path):
    """폰트 F1=Helvetica(보통), F2=Helvetica-Bold를 같은 줄에 섞어 쓰는 최소 PDF."""
    content = (b"BT /F1 18 Tf 40 700 Td (Normal text ) Tj "
               b"/F2 18 Tf (Bold text) Tj "
               b"/F1 18 Tf ( more normal) Tj ET")
    stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R"
        b" /Resources << /Font << /F1 5 0 R /F2 6 0 R >> >> >>",
        stream,
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>",
    ]
    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


def _mini_pdf_italic(path: Path):
    content = b"BT /F1 24 Tf 40 700 Td (Italic text) Tj ET"
    stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R"
        b" /Resources << /Font << /F1 5 0 R >> >> >>",
        stream,
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Oblique >>",
    ]
    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


def _mini_pdf_with_xobject_text(path: Path):
    """텍스트가 Form XObject(벡터 그래픽 등에 흔히 쓰이는 컨테이너) 안에
    그려진 최소 PDF — pdfminer는 페이지 최상위 텍스트와 달리 XObject 내부는
    LTTextLine/LTTextContainer로 자동으로 묶어주지 않는다(재현 확인, 코드
    리뷰 지적: LTTextContainer만 훑으면 이 텍스트를 조용히 놓침)."""
    inner_content = b"BT /F1 18 Tf 0 0 Td (Nested XObject Text) Tj ET"
    xobj_stream = (b"<< /Type /XObject /Subtype /Form /BBox [0 0 300 100] "
                    b"/Resources << /Font << /F1 6 0 R >> >> /Length " +
                    str(len(inner_content)).encode() + b" >>\nstream\n" + inner_content + b"\nendstream")
    page_content = b"q 1 0 0 1 50 700 cm /Fx1 Do Q"
    page_stream = b"<< /Length " + str(len(page_content)).encode() + b" >>\nstream\n" + page_content + b"\nendstream"
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R"
        b" /Resources << /XObject << /Fx1 5 0 R >> >> >>",
        page_stream,
        xobj_stream,
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


def _mini_pdf_mixed_size_same_style(path: Path):
    """굵기/기울임은 그대로인데 글자 크기만 12pt→24pt로 바뀌는 최소 PDF."""
    content = b"BT /F1 12 Tf 40 700 Td (Small size ) Tj /F1 24 Tf (Big size) Tj ET"
    stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R"
        b" /Resources << /Font << /F1 5 0 R >> >> >>",
        stream,
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


class Base(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)


class TestPdfToDocxFormatting(Base):
    def test_bold_run_detected_within_paragraph(self):
        src = self.tmp / "bold.pdf"
        _mini_pdf_with_bold(src)
        out = converters.convert(src, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs]
        by_text = {r.text.strip(): r for r in runs if r.text.strip()}
        self.assertIn("Bold text", by_text)
        self.assertTrue(by_text["Bold text"].font.bold)
        self.assertFalse(by_text["Normal text"].font.bold)

    def test_italic_run_detected(self):
        src = self.tmp / "italic.pdf"
        _mini_pdf_italic(src)
        out = converters.convert(src, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs if r.text.strip()]
        self.assertTrue(runs)
        self.assertTrue(all(r.font.italic for r in runs))

    def test_font_size_preserved(self):
        src = self.tmp / "italic.pdf"
        _mini_pdf_italic(src)
        out = converters.convert(src, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs if r.text.strip()]
        from docx.shared import Pt
        self.assertEqual(runs[0].font.size, Pt(24))

    def test_underline_never_set_pdf_cannot_detect_it(self):
        """DEC-027: PDF는 밑줄을 폰트 속성이 아니라 벡터 선으로 그리는 경우가
        많아 폰트 이름 휴리스틱으로 판별 불가 — 항상 underline=False."""
        src = self.tmp / "bold.pdf"
        _mini_pdf_with_bold(src)
        out = converters.convert(src, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs if r.text.strip()]
        self.assertTrue(all(not r.font.underline for r in runs))

    def test_text_inside_form_xobject_not_lost(self):
        """코드 리뷰 지적: LTTextContainer만 훑는 순회는 Form XObject(벡터
        그래픽 등에 흔히 쓰이는 컨테이너) 안의 텍스트를 조용히 놓쳤다 —
        extract_text() 기반 구경로는 잡아내는데 새 레이아웃 순회만 놓치는
        회귀를 재현 확인 후 재귀 순회로 수정."""
        src = self.tmp / "xobj.pdf"
        _mini_pdf_with_xobject_text(src)
        out = converters.convert(src, "docx", self.tmp)
        texts = [p.text for p in Document(out).paragraphs]
        self.assertTrue(any("Nested XObject Text" in t for t in texts))

    def test_size_change_within_same_bold_italic_splits_run(self):
        """코드 리뷰 지적: flush 조건이 (bold, italic)만 봐서 굵기/기울임이
        그대로인 채 크기만 바뀌면 한 run으로 합쳐지고 마지막 글자 크기로
        전체가 덮였다 — flush 조건에 크기를 포함하도록 수정."""
        src = self.tmp / "mixed_size.pdf"
        _mini_pdf_mixed_size_same_style(src)
        out = converters.convert(src, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs if r.text.strip()]
        from docx.shared import Pt
        by_text = {r.text.strip(): r for r in runs}
        self.assertEqual(by_text["Small size"].font.size, Pt(12))
        self.assertEqual(by_text["Big size"].font.size, Pt(24))


@unittest.skipUnless(HWP_DISTRIBUTION.exists(), "distribution.hwp 샘플 없음(로컬 spike 빌드 필요)")
class TestHwpToDocxFormatting(Base):
    """실제 공공기관 실사용 HWP 문서(distribution.hwp)에 자연히 포함된 굵은
    제목·부분 굵게 강조를 이용 — 합성 픽스처보다 실사용 문서 검증을 우선."""

    def test_title_paragraph_is_bold_and_larger(self):
        out = converters.convert(HWP_DISTRIBUTION, "docx", self.tmp)
        paragraphs = Document(out).paragraphs
        title = next(p for p in paragraphs if "선정 입찰공" in p.text)
        self.assertTrue(all(r.font.bold for r in title.runs if r.text.strip()))

    def test_inline_bold_emphasis_within_body_paragraph(self):
        """"(※부가가치세 포함)"만 굵게 강조된 문단 — 문단 안에서 서식이 바뀌는
        지점마다 run이 정확히 나뉘는지 확인."""
        out = converters.convert(HWP_DISTRIBUTION, "docx", self.tmp)
        paragraphs = Document(out).paragraphs
        body = next(p for p in paragraphs if "부가가치세" in p.text)
        bold_runs = [r for r in body.runs if "부가가치세" in r.text]
        normal_runs = [r for r in body.runs if "기초금액" in r.text]
        self.assertTrue(bold_runs and bold_runs[0].font.bold)
        self.assertTrue(normal_runs and not normal_runs[0].font.bold)

    def test_no_silent_text_loss_from_weighted_position_mismatch(self):
        """코드 리뷰로 발견한 심각한 회귀(머지 후 재검토): run 분리 로직이
        ParaCharShape의 위치를 charList 인덱스와 동일하게 취급했는데, 실제로는
        HWPChar.getCharSize() 가중치 단위(확장/인라인 컨트롤 문자는 charList
        에서 1칸만 차지해도 8로 셈)라 서로 달랐다. 이 문서(distribution.hwp)
        에서 실제로 글자가 조용히 사라지는 걸 재현 확인했다: 단독 특수문자
        "⑤"가 통째로 사라짐(1글자짜리 run이 hwplib 자체의
        ParaText.getNormalString(start,end) 버그—start==end면 무조건 빈
        문자열 반환—에 걸림), 괄호 앞 공백이 2곳에서 사라짐(가중치 위치를
        charList 인덱스로 착각해 문단 뒷부분이 밀림). 손으로 짠 픽스처가
        아니라 실제 공공기관 문서로 검증해야 이런 문제가 잡힌다는 게 이번
        사례의 교훈이라 앞으로도 이 문서를 계속 기준으로 쓴다."""
        out = converters.convert(HWP_DISTRIBUTION, "docx", self.tmp)
        joined = "\n".join(p.text for p in Document(out).paragraphs)
        self.assertIn("⑤전자입찰이용관련", joined)
        self.assertIn("면제함 (각서는", joined)
        self.assertIn("기준」(지방계약법시행령) 및", joined)


@unittest.skipUnless(shutil.which("java"), "JDK 없음")
class TestHwpMakeFormattedFixture(Base):
    """MakeFormattedHwp(테스트 전용 도구, sidecar/hwp/MakeFormattedHwp.java)로
    만든 합성 픽스처 — 실사용 샘플에 없는 기울임·색상·큰 글씨 조합까지 검증."""

    def _classpath(self):
        from app.converters import hwp as hwp_mod
        return hwp_mod._classpath()

    def _java(self):
        from app.converters import hwp as hwp_mod
        return hwp_mod._java()

    def setUp(self):
        super().setUp()
        if self._classpath() is None or self._java() is None:
            self.skipTest("hwplib 클래스패스/JDK 없음(로컬 spike 빌드 필요)")

    def test_bold_italic_size_color_roundtrip_through_docx(self):
        fixture = self.tmp / "formatted.hwp"
        proc = subprocess.run(
            [self._java(), "-cp", self._classpath(), "MakeFormattedHwp", str(fixture)],
            capture_output=True, timeout=30)
        if proc.returncode != 0:
            self.skipTest(f"MakeFormattedHwp 미컴파일 또는 실행 실패: {proc.stderr.decode(errors='replace')}")

        out = converters.convert(fixture, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs]
        by_text = {r.text: r for r in runs}

        self.assertFalse(by_text["일반텍스트"].font.bold)
        self.assertTrue(by_text["굵은텍스트"].font.bold)

        italic_run = by_text["빨간색기울임큰글씨"]
        self.assertTrue(italic_run.font.italic)
        from docx.shared import Pt
        self.assertEqual(italic_run.font.size, Pt(18))
        self.assertEqual(str(italic_run.font.color.rgb), "FF0000")


if __name__ == "__main__":
    unittest.main()
