"""DOCX 표 셀 병합·열 너비 추출/재생성 테스트 (DEC-035).

docx_extract.py(DOCX→블록)·docx_build.py(블록→DOCX)의 순수 파이썬 로직만
검증한다 — HWP 사이드카(Java)는 거치지 않으므로 JDK/hwplib 없이도 항상
실행된다. HWP 왕복까지 포함한 통합 테스트는 test_hwp_table_generation.py에
있다(hwplib 필요, 조건부 스킵).
"""
import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.converters.docx_build import blocks_to_docx
from app.converters.docx_extract import docx_to_blocks


def _cell_text(cell: dict) -> str:
    return "".join(r.get("text") or "" for r in cell["runs"])


class TestDocxExtractMerge(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def _table_docx(self, n_rows, n_cols, merges=()):
        """merges: [(r1, c1, r2, c2), ...] — table.cell(r1,c1)를 (r2,c2)까지 병합."""
        src = self.tmp / "table.docx"
        doc = Document()
        table = doc.add_table(rows=n_rows, cols=n_cols)
        for r in range(n_rows):
            for c in range(n_cols):
                table.cell(r, c).text = f"{r}{c}"
        for r1, c1, r2, c2 in merges:
            table.cell(r1, c1).merge(table.cell(r2, c2))
            table.cell(r1, c1).text = "merged"
        doc.save(src)
        return src

    def test_horizontal_merge_detected(self):
        src = self._table_docx(2, 3, merges=[(0, 0, 0, 1)])
        blocks = docx_to_blocks(src)
        rows = blocks[0]["rows"]
        self.assertEqual(len(rows[0]), 2)  # 병합으로 열이 하나 줄어듦
        self.assertEqual(_cell_text(rows[0][0]), "merged")
        self.assertEqual(rows[0][0]["colSpan"], 2)
        self.assertEqual(rows[0][0]["rowSpan"], 1)
        self.assertEqual(_cell_text(rows[0][1]), "02")
        self.assertEqual(rows[0][1]["colSpan"], 1)
        self.assertEqual(rows[0][1]["rowSpan"], 1)
        self.assertEqual(len(rows[1]), 3)

    def test_vertical_merge_detected(self):
        src = self._table_docx(3, 2, merges=[(0, 1, 1, 1)])
        blocks = docx_to_blocks(src)
        rows = blocks[0]["rows"]
        self.assertEqual(_cell_text(rows[0][1]), "merged")
        self.assertEqual(rows[0][1]["colSpan"], 1)
        self.assertEqual(rows[0][1]["rowSpan"], 2)
        # 세로 병합이 차지한 (1,1)은 두 번째 행에서 생략된다.
        self.assertEqual(len(rows[1]), 1)
        self.assertEqual(_cell_text(rows[1][0]), "10")

    def test_no_merge_matches_old_flat_shape(self):
        src = self._table_docx(2, 2)
        blocks = docx_to_blocks(src)
        rows = blocks[0]["rows"]
        for row in rows:
            for cell in row:
                self.assertEqual(cell["colSpan"], 1)
                self.assertEqual(cell["rowSpan"], 1)

    def test_column_widths_extracted_in_mm(self):
        from docx.shared import Cm

        src = self.tmp / "widths.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(5)
        for c in range(2):
            table.cell(0, c).text = f"c{c}"
        doc.save(src)

        blocks = docx_to_blocks(src)
        widths = blocks[0]["colWidthsMm"]
        self.assertAlmostEqual(widths[0], 30.0, delta=0.1)
        self.assertAlmostEqual(widths[1], 50.0, delta=0.1)


class TestDocxBuildMerge(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def test_string_and_object_cells_both_supported(self):
        """구버전(평문 문자열)·신버전(객체) 셀이 한 표 안에 섞여도 동작해야
        한다 — HwpToJson.java가 병합 안 된 셀은 여전히 평문으로 낸다."""
        blocks = [{"type": "table", "rows": [["plain", {"text": "obj", "colSpan": 1, "rowSpan": 1}]]}]
        out = blocks_to_docx(blocks, self.tmp / "out.docx")
        doc = Document(out)
        table = doc.tables[0]
        self.assertEqual(table.cell(0, 0).text, "plain")
        self.assertEqual(table.cell(0, 1).text, "obj")

    def test_horizontal_merge_rebuilt(self):
        blocks = [{
            "type": "table",
            "rows": [
                [{"text": "merged", "colSpan": 2, "rowSpan": 1}, "02"],
                ["10", "11", "12"],
            ],
        }]
        out = blocks_to_docx(blocks, self.tmp / "out.docx")
        doc = Document(out)
        table = doc.tables[0]
        self.assertEqual(len(table.columns), 3)
        self.assertEqual(table.cell(0, 0).text, "merged")
        # 병합된 두 그리드 칸은 python-docx에서 같은 _tc를 공유한다.
        self.assertEqual(table.cell(0, 0)._tc, table.cell(0, 1)._tc)
        self.assertIn("<w:gridSpan", table._tbl.xml)

    def test_vertical_merge_rebuilt(self):
        blocks = [{
            "type": "table",
            "rows": [
                ["00", {"text": "merged", "colSpan": 1, "rowSpan": 2}],
                ["10"],
            ],
        }]
        out = blocks_to_docx(blocks, self.tmp / "out.docx")
        doc = Document(out)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(table.cell(0, 1)._tc, table.cell(1, 1)._tc)
        self.assertIn("<w:vMerge", table._tbl.xml)

    def test_extract_then_build_round_trip(self):
        """docx_extract가 만든 블록을 docx_build에 그대로 먹여도 병합이
        살아남는지 — 두 모듈이 같은 표현을 공유함을 직접 확인."""
        from app.converters.docx_extract import docx_to_blocks

        src = self.tmp / "src.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        for r in range(2):
            for c in range(2):
                table.cell(r, c).text = f"{r}{c}"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = "merged"
        doc.save(src)

        blocks = docx_to_blocks(src)
        out = blocks_to_docx(blocks, self.tmp / "rebuilt.docx")
        rebuilt = Document(out)
        table2 = rebuilt.tables[0]
        self.assertEqual(table2.cell(0, 0).text, "merged")
        self.assertEqual(table2.cell(0, 0)._tc, table2.cell(0, 1)._tc)

    def test_cell_char_formatting_round_trip(self):
        """표 셀 안 문자 서식 보존 개선 — 굵게/기울임/색상이 셀의 run
        단위로 그대로 추출·재생성되는지(hwplib/hwpxlib을 거치지 않는
        순수 파이썬 경로만 검증, 사이드카 왕복은 test_hwp_table_generation.py/
        test_hwpx.py에 있음)."""
        from docx.shared import RGBColor
        from app.converters.docx_extract import docx_to_blocks

        src = self.tmp / "formatted.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        p = table.cell(0, 0).paragraphs[0]
        p.add_run("일반 ")
        bold_run = p.add_run("굵게빨강")
        bold_run.bold = True
        bold_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        doc.save(src)

        blocks = docx_to_blocks(src)
        cell = blocks[0]["rows"][0][0]
        runs = cell["runs"]
        self.assertEqual([r["text"] for r in runs], ["일반 ", "굵게빨강"])
        self.assertFalse(runs[0]["bold"])
        self.assertTrue(runs[1]["bold"])
        self.assertEqual(runs[1]["color"], "FF0000")

        out = blocks_to_docx(blocks, self.tmp / "rebuilt2.docx")
        rebuilt = Document(out)
        rebuilt_runs = rebuilt.tables[0].cell(0, 0).paragraphs[0].runs
        self.assertEqual([r.text for r in rebuilt_runs], ["일반 ", "굵게빨강"])
        self.assertFalse(bool(rebuilt_runs[0].font.bold))
        self.assertTrue(rebuilt_runs[1].font.bold)
        self.assertEqual(rebuilt_runs[1].font.color.rgb, RGBColor(0xFF, 0x00, 0x00))

    def test_cell_alignment_round_trip(self):
        """표 셀 정렬 보존 개선 — 셀 안 문단 정렬(가운데/오른쪽)이 그대로
        추출·재생성되는지(hwplib/hwpxlib을 거치지 않는 순수 파이썬 경로만
        검증, 사이드카 왕복은 test_hwp_table_generation.py/test_hwpx.py에
        있음)."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from app.converters.docx_extract import docx_to_blocks

        src = self.tmp / "aligned.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 0).paragraphs[0].add_run("가운데")
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        table.cell(0, 1).paragraphs[0].add_run("오른쪽")
        doc.save(src)

        blocks = docx_to_blocks(src)
        row = blocks[0]["rows"][0]
        self.assertEqual(row[0]["align"], "center")
        self.assertEqual(row[1]["align"], "right")

        out = blocks_to_docx(blocks, self.tmp / "rebuilt3.docx")
        rebuilt = Document(out)
        rebuilt_cells = rebuilt.tables[0].rows[0].cells
        self.assertEqual(rebuilt_cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(rebuilt_cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)


if __name__ == "__main__":
    unittest.main()
