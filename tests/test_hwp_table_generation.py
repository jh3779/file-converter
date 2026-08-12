"""DOCX→HWP 표 신규 생성 테스트 — DEC-028 (DEC-017 정정).

hwplib이 실제로는 표를 처음부터 만들 수 있다는 게 확인되면서(공식 샘플
src/test/sample/Inserting_Table.java), 이전엔 " | "로 이어붙인 텍스트로
단순화하던 DOCX 표를 이제 실제 HWP 표 컨트롤로 생성한다
(spike/hwplib/SpikeTable.java에서 왕복 검증 후 sidecar/hwp/JsonToHwp.java에
일반화). 실제 hwplib 라이브러리가 필요해 기존 TestHwp와 동일한 조건에서만
실행 — 로컬 spike 빌드가 없는 CI "test" job에서는 스킵된다(실제 CI 게이트는
build.yml의 HWP 엔진 스모크 쪽에 별도로 있음).

**중요한 한계**: 이 테스트들은 hwplib 자체의 왕복 읽기(HwpToJson)로만
검증한다 — 실제 한글/한워드 뷰어에서 표가 정확히 렌더링되는지는 이
테스트로 확인할 수 없다(DEC-018과 동일한 근본적 제약, Mac 개발 환경에는
뷰어가 없음). Windows 실사용자 테스트가 최종 검증에 반드시 필요하다.
"""
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document

from app import converters

REPO = Path(__file__).resolve().parents[1]
HWP_SAMPLE = REPO / "spike" / "hwplib" / "repo" / "sample_hwp" / "basic" / "표.hwp"


def _hwp_available():
    from app.converters import hwp as hwp_mod
    return hwp_mod._java() is not None and hwp_mod._classpath() is not None


@unittest.skipUnless(
    HWP_SAMPLE.exists() and shutil.which("java") and _hwp_available(),
    "hwplib 샘플/JDK 없음 — spike 빌드 후 실행 (RESULT.md)")
class TestDocxToHwpTableGeneration(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _docx_with_table(self, rows_data):
        src = self.tmp / "table.docx"
        doc = Document()
        n_rows = len(rows_data)
        n_cols = len(rows_data[0])
        table = doc.add_table(rows=n_rows, cols=n_cols)
        for r, row in enumerate(rows_data):
            for c, text in enumerate(row):
                table.cell(r, c).text = text
        doc.save(src)
        return src

    def test_table_survives_as_real_table_not_flattened_text(self):
        """DEC-028의 핵심: 결과 HWP를 다시 DOCX로 읽었을 때 실제 표
        (doc.tables)로 나와야 한다 — 이전처럼 " | "로 이어붙은 문단이
        아니라."""
        src = self._docx_with_table([["이름", "부서"], ["김철수", "영업1팀"]])
        out = converters.convert(src, "hwp", self.tmp)
        self.assertTrue(out.exists())

        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc = Document(back)

        self.assertTrue(doc.tables, "표가 문단으로 단순화됨(DEC-028 회귀) — doc.tables가 비어 있음")
        cells = [c.text for row in doc.tables[0].rows for c in row.cells]
        self.assertIn("이름", cells)
        self.assertIn("김철수", cells)
        self.assertIn("영업1팀", cells)
        # " | "로 이어붙인 옛 방식의 흔적이 문단에 남아있지 않아야 한다.
        joined_paragraphs = " ".join(p.text for p in doc.paragraphs)
        self.assertNotIn("김철수 | 영업1팀", joined_paragraphs)

    def test_table_row_and_column_count_preserved(self):
        src = self._docx_with_table([
            ["A", "B", "C"],
            ["1", "2", "3"],
            ["4", "5", "6"],
        ])
        out = converters.convert(src, "hwp", self.tmp)
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc = Document(back)

        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 3)

    def test_paragraphs_before_and_after_table_preserved_in_order(self):
        """문단 → 표 → 문단처럼 섞인 문서에서 순서와 내용이 모두 보존되는지."""
        src = self.tmp / "mixed.docx"
        doc = Document()
        doc.add_paragraph("직원 명단입니다")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "이름"
        table.cell(0, 1).text = "부서"
        doc.add_paragraph("이상입니다")
        doc.save(src)

        out = converters.convert(src, "hwp", self.tmp)
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        texts = [p.text for p in doc2.paragraphs if p.text.strip()]
        self.assertIn("직원 명단입니다", texts)
        self.assertIn("이상입니다", texts)
        self.assertTrue(doc2.tables)
        self.assertEqual([c.text for c in doc2.tables[0].rows[0].cells], ["이름", "부서"])

    def test_merged_cells_survive_docx_to_hwp_to_docx_round_trip(self):
        """DEC-035: DOCX의 병합된 셀(가로+세로)이 HWP를 거쳐 다시 DOCX로
        와도 실제 병합(gridSpan/vMerge)으로 남아있어야 한다 — 이전에는
        병합이 풀려 평문 그리드로만 나왔다(DEC-028 "알려진 한계")."""
        src = self.tmp / "merged.docx"
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        for r in range(3):
            for c in range(3):
                table.cell(r, c).text = f"{r}{c}"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = "H-merged"
        table.cell(1, 2).merge(table.cell(2, 2))
        table.cell(1, 2).text = "V-merged"
        doc.save(src)

        out = converters.convert(src, "hwp", self.tmp)
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        table2 = doc2.tables[0]
        self.assertEqual(len(table2.rows), 3)
        self.assertEqual(len(table2.columns), 3)
        self.assertEqual(table2.cell(0, 0).text, "H-merged")
        self.assertEqual(table2.cell(0, 0)._tc, table2.cell(0, 1)._tc,
                          "가로 병합이 풀려서 돌아옴")
        self.assertEqual(table2.cell(1, 2).text, "V-merged")
        self.assertEqual(table2.cell(1, 2)._tc, table2.cell(2, 2)._tc,
                          "세로 병합이 풀려서 돌아옴")

    def test_cell_char_formatting_survives_docx_to_hwp_to_docx_round_trip(self):
        """표 셀 안 문자 서식 보존 개선 — 굵게·기울임·밑줄·크기·색상이
        표 셀의 run 단위로 반영되고, HWP를 거쳐 다시 DOCX로 와도
        그대로 남아있어야 한다(DEC-038이 문단에 반영한 것과 대칭,
        DEC-051)."""
        from docx.shared import Pt, RGBColor

        src = self.tmp / "cell_formatted.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        p = table.cell(0, 0).paragraphs[0]
        p.add_run("일반 ")
        styled = p.add_run("굵고빨간18pt")
        styled.bold = True
        styled.font.size = Pt(18)
        styled.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        table.cell(0, 1).text = "plain"
        doc.save(src)

        out = converters.convert(src, "hwp", self.tmp)
        back_dir = self.tmp / "back_fmt"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        cell0_runs = doc2.tables[0].cell(0, 0).paragraphs[0].runs
        by_text = {r.text: r for r in cell0_runs}
        self.assertFalse(bool(by_text["일반 "].font.bold))
        self.assertTrue(by_text["굵고빨간18pt"].font.bold)
        self.assertEqual(by_text["굵고빨간18pt"].font.size, Pt(18))
        self.assertEqual(by_text["굵고빨간18pt"].font.color.rgb, RGBColor(0xFF, 0x00, 0x00))
        self.assertEqual(doc2.tables[0].cell(0, 1).text, "plain")

    def test_vertical_merge_spanning_entire_row_preserves_row_count(self):
        """세로 병합이 어떤 행 전체를 덮으면(예: 1열 표에서 위 셀이 아래
        행까지 병합) docx_extract.py는 그 행을 빈 배열([])로 낸다 — 이
        빈 행을 JsonToHwp.parseTableSpec이 버리면 rowCount가 원본보다
        줄어 뒤 행들의 병합·좌표가 전부 어긋난다(자동 PR 리뷰로 발견해
        수정). 3행×1열, 위 두 행이 하나로 병합, 마지막 행은 별개인 표로
        재현한다."""
        src = self.tmp / "full_row_merge.docx"
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        table.cell(0, 0).text = "머리글"
        table.cell(1, 0).text = ""
        table.cell(2, 0).text = "마지막행"
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "머리글"
        doc.save(src)

        out = converters.convert(src, "hwp", self.tmp)
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        table2 = doc2.tables[0]
        self.assertEqual(len(table2.rows), 3, "행 전체를 덮는 세로 병합 때문에 행 수가 줄어듦(회귀)")
        self.assertEqual(table2.cell(0, 0).text, "머리글")
        self.assertEqual(table2.cell(0, 0)._tc, table2.cell(1, 0)._tc,
                          "세로 병합이 풀려서 돌아옴")
        self.assertEqual(table2.cell(2, 0).text, "마지막행")

    def test_hwp_to_json_structure_matches_source_dimensions(self):
        """HwpToJson으로 재읽었을 때도 실제 표 블록(행렬 수 일치)으로
        나오는지 — Python DOCX 재변환 경로를 거치지 않는 더 직접적인 확인.
        DEC-051부터 셀은 항상 {"runs":[...],"colSpan":n,"rowSpan":m} 객체다."""
        import json
        from app.converters import hwp as hwp_mod

        src = self._docx_with_table([["x", "y", "z"], ["1", "2", "3"]])
        out = converters.convert(src, "hwp", self.tmp)

        blocks_json = self.tmp / "readback.json"
        hwp_mod._run_sidecar("HwpToJson", out, blocks_json)
        data = json.loads(blocks_json.read_text(encoding="utf-8"))
        table_blocks = [b for b in data["blocks"] if b["type"] == "table"]
        self.assertEqual(len(table_blocks), 1)
        rows = table_blocks[0]["rows"]
        self.assertEqual(len(rows), 2)
        self.assertEqual(len(rows[0]), 3)
        texts = [["".join(r["text"] for r in cell["runs"]) for cell in row] for row in rows]
        self.assertEqual(texts, [["x", "y", "z"], ["1", "2", "3"]])


if __name__ == "__main__":
    unittest.main()
