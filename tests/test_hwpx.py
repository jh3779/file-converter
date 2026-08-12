"""HWPX 변환 테스트 — QA(h) Phase 1(읽기, 외부 QA 요청).

hwpxlib(spike/hwpxlib/RESULT.md) 로컬 빌드가 있을 때만 실행된다(사이드카가
필요 없는 다른 테스트들과 달리 실제 JRE + hwpxlib 클래스로 sidecar를
실행해야 함 — tests/test_pipeline.py의 TestHwp와 같은 조건).
hwpxlib 저장소 자체의 테스트 픽스처(testFile/tool/textextractor/*.hwpx)를
그대로 쓴다 — 이 프로젝트 저장소에는 hwpx 샘플을 직접 넣지 않는다(제3자
문서를 저장소에 넣지 않는다는 기존 관례, DEC-032와 같은 원칙).
"""
import shutil
import subprocess
import tempfile
import unittest
from pathlib import Path

from app import converters

REPO = Path(__file__).resolve().parents[1]
HWPX_SAMPLE = REPO / "spike" / "hwpxlib" / "repo" / "testFile" / "tool" / "textextractor" / "multipara.hwpx"
HWPX_TABLE_SAMPLE = REPO / "spike" / "hwpxlib" / "repo" / "testFile" / "tool" / "textextractor" / "Table.hwpx"
HWPX_HEADER_FOOTER_SAMPLE = REPO / "spike" / "hwpxlib" / "repo" / "testFile" / "reader_writer" / "HeaderFooter.hwpx"
HWPX_NESTED_SHAPE_SAMPLE = REPO / "spike" / "hwpxlib" / "repo" / "testFile" / "tool" / "textextractor" / "RectInRect.hwpx"


def _find_soffice():
    from app.converters import office
    return office.find_soffice()


def _make_tab_hwpx(out_path: Path):
    """MakeTabHwpx(테스트 전용 디버그 도구, sidecar/hwp/MakeTabHwpx.java)를
    실행해 "가나"+Tab+"다라" 문단 하나짜리 최소 HWPX를 만든다 — test_pipeline.py의
    _run_linesegdebug와 같은 패턴(로컬 build.sh 산출물을 subprocess로 직접 실행)."""
    from app.converters import hwp as hwp_mod
    java = hwp_mod._java()
    cp = hwp_mod._classpath()
    proc = subprocess.run([java, "-cp", cp, "MakeTabHwpx", str(out_path)],
                           capture_output=True, text=True, timeout=30)
    assert proc.returncode == 0, proc.stderr


class Base(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)


@unittest.skipUnless(
    HWPX_SAMPLE.exists() and shutil.which("java"),
    "hwpxlib 샘플/JDK 없음 — spike/hwpxlib/RESULT.md 절차로 로컬 빌드 후 실행")
class TestHwpx(Base):
    def test_hwpx_to_txt(self):
        out = converters.convert(HWPX_SAMPLE, "txt", self.tmp)
        text = out.read_text(encoding="utf-8")
        self.assertIn("김하성", text)

    def test_hwpx_to_docx_preserves_paragraphs(self):
        from docx import Document
        out = converters.convert(HWPX_SAMPLE, "docx", self.tmp)
        texts = [p.text for p in Document(out).paragraphs]
        self.assertTrue(any("김하성" in t for t in texts))
        self.assertGreater(len(texts), 1)  # 여러 문단으로 나뉘어 있어야 함(한 덩어리로 뭉개지면 안 됨)

    def test_hwpx_to_docx_preserves_table(self):
        from docx import Document
        out = converters.convert(HWPX_TABLE_SAMPLE, "docx", self.tmp)
        tables = Document(out).tables
        self.assertTrue(tables)
        rows = [[c.text for c in row.cells] for row in tables[0].rows]
        self.assertEqual(rows[0], ["이름", "국어", "영어", "수학"])
        self.assertIn(["개똥이", "89", "65", "78"], rows)

    def test_hwpx_to_docx_preserves_char_formatting(self):
        """Table.hwpx의 "날짜" 문단은 기울임+빨간색으로 서식이 지정돼 있다
        (로컬 조사로 확인, spike/hwpxlib/RESULT.md) — HwpxToJson의 CharPr
        추출이 실제로 동작하는지 회귀 확인."""
        from docx import Document
        out = converters.convert(HWPX_TABLE_SAMPLE, "docx", self.tmp)
        runs = [r for p in Document(out).paragraphs for r in p.runs if r.text.strip()]
        by_text = {r.text.strip(): r for r in runs}
        self.assertIn("날짜", by_text)
        self.assertTrue(by_text["날짜"].font.italic)

    def test_hwpx_to_docx_preserves_header_footer_text(self):
        """HWPX Phase 1(DEC-044)·Phase 2(DEC-049)가 반복적으로 "범위 밖"으로
        문서화해온 한계 — 머리말/꼬리말이 Ctrl RunItem으로 감싸여 있어
        기존 emitParagraph가 아예 순회하지 않고 조용히 건너뛰었다(HwpToJson.java
        쪽이 외부 QA #43로 이미 겪은 것과 같은 종류의 문제, DEC-032). 재귀
        처리 추가 확인용 — HeaderFooter.hwpx는 본문이 없고 머리말/꼬리말
        텍스트만 있는 픽스처라, 수정 전에는 결과 DOCX 문단이 전부 비어
        있었다(직접 재현 확인)."""
        from docx import Document
        out = converters.convert(HWPX_HEADER_FOOTER_SAMPLE, "docx", self.tmp)
        texts = [p.text for p in Document(out).paragraphs if p.text.strip()]
        self.assertIn("머리말 테스트", texts)
        self.assertIn("꼬리말", texts)

    def test_hwpx_to_docx_preserves_nested_shape_text(self):
        """글상자(도형) 안, 그리고 도형을 묶은 그룹(Container) 안의 텍스트도
        재귀로 뽑아내는지 — RectInRect.hwpx는 사각형 안에 사각형이 중첩된
        픽스처라, 수정 전에는 결과 DOCX 문단이 전부 비어 있었다(직접 재현
        확인, hwplib 쪽 ControlContainer 중첩 그룹 처리와 대칭)."""
        from docx import Document
        out = converters.convert(HWPX_NESTED_SHAPE_SAMPLE, "docx", self.tmp)
        texts = [p.text for p in Document(out).paragraphs if p.text.strip()]
        self.assertTrue(texts, "글상자 안 텍스트가 하나도 안 나옴(회귀)")
        self.assertTrue(any("사각형" in t for t in texts))

    def test_hwpx_tab_normalized_to_space_not_dropped(self):
        """HwpxToJson(HWPX→DOCX/PDF 경로, HWPX→TXT는 hwpxlib 자체
        TextExtractor를 써서 별도)의 extractTextFrom이 Tab(TItem)을 조용히
        건너뛰면 탭 앞뒤 텍스트가 공백 없이 그대로 붙어버리는 회귀가
        있었다(수정 확인용) — "가나"+Tab+"다라"가 "가나다라"로 뭉개지지
        않고 "가나 다라"로 나와야 한다."""
        from docx import Document
        src = self.tmp / "tab.hwpx"
        _make_tab_hwpx(src)
        out = converters.convert(src, "docx", self.tmp)
        texts = [p.text for p in Document(out).paragraphs]
        self.assertIn("가나 다라", texts)
        self.assertNotIn("가나다라", texts)

    @unittest.skipUnless(_find_soffice(), "LibreOffice(soffice) 없음 — 로컬/Windows CI에서만 실행")
    def test_hwpx_to_pdf(self):
        out = converters.convert(HWPX_SAMPLE, "pdf", self.tmp)
        self.assertEqual(out.read_bytes()[:5], b"%PDF-")


def _run_pagebreakdebug_hwpx(hwpx_path: Path):
    """PageBreakDebugHwpx(테스트 전용 디버그 도구, sidecar/hwp/
    PageBreakDebugHwpx.java)를 직접 실행해 문단별 (pageBreakBefore, 텍스트)를
    돌려준다 — DEC-049(DEC-039 대칭). HwpxToJson 정식 스키마는 이 속성을
    안 보므로(HWP 쪽과 동일한 이유) hwpxlib으로 직접 열어 확인해야 한다."""
    from app.converters import hwp as hwp_mod
    java = hwp_mod._java()
    cp = hwp_mod._classpath()
    with tempfile.TemporaryDirectory() as tmp:
        out_path = Path(tmp) / "pagebreak-debug.txt"
        proc = subprocess.run([java, "-cp", cp, "PageBreakDebugHwpx", str(hwpx_path), str(out_path)],
                               capture_output=True, text=True, timeout=30)
        assert proc.returncode == 0, proc.stderr
        content = out_path.read_text(encoding="utf-8")
    rows = []
    for line in content.strip().splitlines():
        idx, page_break_before, text = line.split("\t", 2)
        rows.append((int(idx), page_break_before == "true", text))
    return rows


@unittest.skipUnless(
    HWPX_SAMPLE.exists() and shutil.which("java"),
    "hwpxlib 샘플/JDK 없음 — spike/hwpxlib/RESULT.md 절차로 로컬 빌드 후 실행")
class TestHwpxWrite(Base):
    """DOCX/PDF → HWPX 쓰기(HWPX Phase 2, DEC-049) — hwp.py의 docx_to_hwp/
    pdf_to_hwp 왕복 테스트(tests/test_pipeline.py TestHwp)와 대칭."""

    def test_docx_to_hwpx_roundtrip(self):
        from docx import Document
        src = self.tmp / "한글.docx"
        doc = Document()
        doc.add_paragraph("뷁 밟 닳 넋 앎 옳 훑 흙 삵 값 넓 얹 앉 닭 없")
        doc.add_paragraph("大韓民國 韓國語 漢字 契約書 委任狀")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "이름"
        table.cell(0, 1).text = "부서"
        table.cell(1, 0).text = "김철수"
        table.cell(1, 1).text = "영업1팀"
        doc.save(src)

        out = converters.convert(src, "hwpx", self.tmp)
        self.assertTrue(out.exists())

        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)
        texts = [p.text for p in doc2.paragraphs]
        self.assertIn("뷁 밟 닳 넋 앎 옳 훑 흙 삵 값 넓 얹 앉 닭 없", texts)
        self.assertIn("大韓民國 韓國語 漢字 契約書 委任狀", texts)
        self.assertTrue(doc2.tables)
        cells = [c.text for row in doc2.tables[0].rows for c in row.cells]
        self.assertIn("이름", cells)
        self.assertIn("김철수", cells)
        self.assertIn("영업1팀", cells)

    def test_docx_to_hwpx_char_formatting_roundtrip(self):
        """DEC-049(DEC-038 대칭): 문단 문자 서식(굵게/기울임/밑줄/크기/색상)이
        DOCX→HWPX 쓰기 방향에도 반영되는지 — 실제 hwpxlib으로
        DOCX→HWPX→DOCX 왕복 확인."""
        from docx import Document
        from docx.shared import Pt, RGBColor

        src = self.tmp / "formatted.docx"
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("일반 ")
        bold_run = p.add_run("굵게")
        bold_run.bold = True
        styled_run = p.add_run("기울임+18pt+빨강")
        styled_run.italic = True
        styled_run.font.size = Pt(18)
        styled_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        p2 = doc.add_paragraph()
        p2.add_run("둘째 문단 ")
        underline_run = p2.add_run("밑줄")
        underline_run.underline = True
        doc.save(src)

        out = converters.convert(src, "hwpx", self.tmp)
        back_dir = self.tmp / "back_fmt"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        by_text = {run.text: run for p in doc2.paragraphs for run in p.runs}
        self.assertFalse(bool(by_text["일반 "].font.bold))
        self.assertTrue(by_text["굵게"].font.bold)
        self.assertTrue(by_text["기울임+18pt+빨강"].font.italic)
        self.assertEqual(by_text["기울임+18pt+빨강"].font.size, Pt(18))
        self.assertEqual(by_text["기울임+18pt+빨강"].font.color.rgb, RGBColor(0xFF, 0x00, 0x00))
        self.assertTrue(by_text["밑줄"].font.underline)

    def test_docx_to_hwpx_table_merge_roundtrip(self):
        """DEC-049(DEC-035 대칭): DOCX의 병합된 셀(가로+세로)이 HWPX를
        거쳐 다시 DOCX로 와도 실제 병합(gridSpan/vMerge)으로 남아있어야
        한다 — hwpxlib의 sparse 표현(스파이크로 확인, spike/hwpxlib/
        RESULT.md "Phase 2(쓰기)")이 정확히 왕복하는지 검증."""
        from docx import Document
        src = self.tmp / "merged.docx"
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        for r in range(3):
            for c in range(3):
                table.cell(r, c).text = f"{r}{c}"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = "H-merged"
        table.cell(1, 2).merge(table.cell(2, 2))
        table.cell(1, 2).text = "V-merged"
        doc.save(src)

        out = converters.convert(src, "hwpx", self.tmp)
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        table2 = doc2.tables[0]
        self.assertEqual(len(table2.rows), 3)
        self.assertEqual(len(table2.columns), 3)
        self.assertEqual(table2.cell(0, 0).text, "H-merged")
        self.assertEqual(table2.cell(0, 0)._tc, table2.cell(0, 1)._tc,
                          "가로 병합이 풀려서 돌아옴")
        self.assertEqual(table2.cell(1, 2).text, "V-merged")
        self.assertEqual(table2.cell(1, 2)._tc, table2.cell(2, 2)._tc,
                          "세로 병합이 풀려서 돌아옴")

    def test_docx_to_hwpx_cell_char_formatting_roundtrip(self):
        """표 셀 안 문자 서식 보존 개선(DEC-051) — HWP 쪽과 대칭. 굵게·
        크기·색상이 표 셀의 run 단위로 반영되고, HWPX를 거쳐 다시 DOCX로
        와도 그대로 남아있어야 한다."""
        from docx import Document
        from docx.shared import Pt, RGBColor

        src = self.tmp / "cell_formatted.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        p = table.cell(0, 0).paragraphs[0]
        p.add_run("일반 ")
        styled = p.add_run("굵고빨간18pt")
        styled.bold = True
        styled.font.size = Pt(18)
        styled.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        table.cell(0, 1).text = "plain"
        doc.save(src)

        out = converters.convert(src, "hwpx", self.tmp)
        back_dir = self.tmp / "back_cell_fmt"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        doc2 = Document(back)

        cell0_runs = doc2.tables[0].cell(0, 0).paragraphs[0].runs
        by_text = {r.text: r for r in cell0_runs}
        self.assertFalse(bool(by_text["일반 "].font.bold))
        self.assertTrue(by_text["굵고빨간18pt"].font.bold)
        self.assertEqual(by_text["굵고빨간18pt"].font.size, Pt(18))
        self.assertEqual(by_text["굵고빨간18pt"].font.color.rgb, RGBColor(0xFF, 0x00, 0x00))
        self.assertEqual(doc2.tables[0].cell(0, 1).text, "plain")

    def test_docx_to_hwpx_alignment_roundtrip(self):
        """DEC-049(DEC-040 대칭): DOCX 문단에 직접 지정된 정렬이 HWPX의
        실제 ParaPr.align()으로 반영되고, 다시 DOCX로 왕복해도 남는지."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        src = self.tmp / "정렬.docx"
        doc = Document()
        doc.add_paragraph("기본 정렬")
        p_center = doc.add_paragraph("가운데 정렬")
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right = doc.add_paragraph("오른쪽 정렬")
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save(src)

        out = converters.convert(src, "hwpx", self.tmp)
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "docx", back_dir)
        by_text = {p.text: p.alignment for p in Document(back).paragraphs}

        self.assertEqual(by_text["기본 정렬"], WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(by_text["가운데 정렬"], WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(by_text["오른쪽 정렬"], WD_ALIGN_PARAGRAPH.RIGHT)

    def test_json_to_hwpx_accepts_plain_string_table_cells(self):
        """DEC-049 머지 전 자동 리뷰 지적(재현 확인 후 수정): JsonToHwpx의
        parseTableSpec()이 셀을 항상 Map으로 캐스팅해, HwpxToJson이 병합
        없는(1×1) 셀을 평문 문자열로 내보내는 것(하위 호환 원칙, DEC-035와
        동일)과 왕복이 안 됐다(ClassCastException) — 실제 앱 파이프라인은
        docx_extract.py/pdf.py가 항상 객체 셀만 내므로 이 경로를 타지
        않지만, HwpxToJson 출력을 그대로 다시 JsonToHwpx에 먹이는 왕복
        도구로서는 깨져 있었다. 문자열 셀도 {text, colSpan:1, rowSpan:1}로
        정규화하도록 수정 확인."""
        import json
        from app.converters.hwp import _run_sidecar

        blocks_json = self.tmp / "string-cells.blocks.json"
        blocks_json.write_text(
            json.dumps({"blocks": [{"type": "table", "rows": [["a", "b"], ["c", "d"]]}]}),
            encoding="utf-8")
        out = self.tmp / "string-cells.hwpx"
        _run_sidecar("JsonToHwpx", blocks_json, out)
        self.assertTrue(out.exists())

        reread_json = self.tmp / "string-cells-reread.json"
        _run_sidecar("HwpxToJson", out, reread_json)
        rows = json.loads(reread_json.read_text(encoding="utf-8"))["blocks"][0]["rows"]
        texts = [["".join(r["text"] for r in cell["runs"]) for cell in row] for row in rows]
        self.assertEqual(texts, [["a", "b"], ["c", "d"]])

    def test_pdf_to_hwpx_preserves_page_breaks(self):
        """DEC-049(DEC-039 대칭): PDF 여러 페이지가 HWPX 안에서도 페이지
        구분 없이 하나로 이어 붙지 않아야 한다 — 각 페이지 첫 문단에
        pageBreakBefore가 반영되는지 PageBreakDebugHwpx로 직접 확인."""
        import sys
        sys.path.insert(0, str(Path(__file__).resolve().parent))
        from test_pipeline import _mini_pdf_pages

        src = self.tmp / "pages.pdf"
        _mini_pdf_pages(src, ["Page One Text", "Page Two Text", "Page Three Text"])
        out = converters.convert(src, "hwpx", self.tmp)

        rows = _run_pagebreakdebug_hwpx(out)
        by_text = {text: page_break_before for _, page_break_before, text in rows}
        self.assertEqual(by_text.get("Page One Text"), False)
        self.assertEqual(by_text.get("Page Two Text"), True)
        self.assertEqual(by_text.get("Page Three Text"), True)

        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = converters.convert(out, "txt", back_dir)
        back_text = back.read_text(encoding="utf-8")
        self.assertIn("Page One Text", back_text)
        self.assertIn("Page Two Text", back_text)
        self.assertIn("Page Three Text", back_text)
        self.assertLess(back_text.index("Page One Text"), back_text.index("Page Two Text"))
        self.assertLess(back_text.index("Page Two Text"), back_text.index("Page Three Text"))


if __name__ == "__main__":
    unittest.main()
