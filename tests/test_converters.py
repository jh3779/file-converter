"""데이터 변환기·출력 규칙 테스트: python -m unittest discover tests"""
import csv
import json
import shutil
import tempfile
import unittest
from pathlib import Path

from app.converters import data
from app.converters import pdf as pdf_mod
from app.converters.base import ConversionError
from app.converters.docx_build import EAST_ASIAN_FONT, blocks_to_docx
from app.converters.docx_extract import docx_to_blocks
from app.converters.pdf import _classify_alignment
from app.output import unique_output_path


class Base(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)


def _mini_pdf_pages(path: Path, texts: list[str]):
    """tests/test_pipeline.py의 _mini_pdf_pages와 같은 원리(최소 PDF를
    페이지마다 다른 텍스트로 생성) — JDK 없이 pdf.py의 순수 파이썬 추출
    로직만 검증하려고 이 파일에도 둔다(다른 테스트 파일들도 각자 최소 PDF
    헬퍼를 따로 둠, DRY보다 파일별 독립성 우선)."""
    n = len(texts)
    page_obj_start = 3
    font_obj_num = page_obj_start + 2 * n
    kids, page_bodies, content_bodies = [], [], []
    for i, text in enumerate(texts):
        page_idx = page_obj_start + 2 * i
        content_idx = page_idx + 1
        kids.append(f"{page_idx} 0 R")
        content = f"BT /F1 18 Tf 40 700 Td ({text}) Tj ET".encode()
        stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"
        page_bodies.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents {content_idx} 0 R"
            f" /Resources << /Font << /F1 {font_obj_num} 0 R >> >> >>".encode())
        content_bodies.append(stream)
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        ("<< /Type /Pages /Kids [" + " ".join(kids) + f"] /Count {n} >>").encode(),
    ]
    for pb, cb in zip(page_bodies, content_bodies):
        objs.append(pb)
        objs.append(cb)
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


class TestPdfToHwpPageBreaks(Base):
    """DEC-039: pdf.py._extract_pdf_blocks_by_page의 순수 파이썬 로직
    (JDK 불요) — 실제 hwplib 왕복 검증은 tests/test_pipeline.py::TestHwp::
    test_pdf_to_hwp_preserves_page_breaks(hwplib 로컬 빌드 있을 때만)."""

    def test_first_page_first_block_has_no_page_break(self):
        pdf = self.tmp / "one.pdf"
        _mini_pdf_pages(pdf, ["Only Page"])
        blocks = pdf_mod._extract_pdf_blocks_by_page(pdf)
        self.assertEqual(len(blocks), 1)
        self.assertNotIn("pageBreakBefore", blocks[0])
        self.assertEqual(blocks[0]["text"], "Only Page")

    def test_later_pages_first_block_marked(self):
        pdf = self.tmp / "three.pdf"
        _mini_pdf_pages(pdf, ["Page One", "Page Two", "Page Three"])
        blocks = pdf_mod._extract_pdf_blocks_by_page(pdf)
        self.assertEqual([b["text"] for b in blocks], ["Page One", "Page Two", "Page Three"])
        self.assertNotIn("pageBreakBefore", blocks[0])
        self.assertTrue(blocks[1]["pageBreakBefore"])
        self.assertTrue(blocks[2]["pageBreakBefore"])

    def test_whitespace_only_first_container_does_not_consume_page_break(self):
        """JsonToHwp.java는 text.trim().isEmpty()인 문단을 버린다 — 페이지
        첫 컨테이너가 공백뿐이면 그 문단이 pageBreakBefore를 달고 있어도
        JsonToHwp 쪽에서 통째로 버려져, 실제 첫(비공백) 문단에는 쪽 나눔이
        반영되지 않는 회귀가 있었다(자동 리뷰로 발견). 공백뿐인 컨테이너는
        first_on_page를 소비하지 않아야 한다.

        실제 PDF로 재현하면 pdfminer의 레이아웃 그룹핑이 컨테이너 순서를
        보장하지 않아(공백 컨테이너가 먼저 온다고 확신할 수 없음) 결과가
        들쭉날쭉해진다 — _paragraph_candidates를 모킹해 "페이지 첫 컨테이너가
        공백뿐"인 순서를 직접 통제한다."""
        from unittest.mock import patch

        # str 서브클래스라 "page0"과의 동일성 비교(side_effect)는 그대로 되면서
        # page.width(정렬 판정용, DEC-040)도 갖는 가짜 페이지 객체.
        class _FakePage(str):
            width = 612.0

        page0, page1 = _FakePage("page0"), _FakePage("page1")

        # extract_pages는 _extract_pdf_blocks_by_page 안에서 지역 import되므로
        # pdf_mod 네임스페이스가 아니라 pdfminer.high_level 자체를 패치해야 한다.
        with patch("pdfminer.high_level.extract_pages", return_value=[page0, page1]), \
             patch.object(pdf_mod, "_paragraph_candidates",
                           side_effect=lambda page: ["c1"] if page == "page0" else ["ws", "c2"]), \
             patch.object(pdf_mod, "_container_to_runs",
                           side_effect=lambda c: [{"text": {"c1": "First page", "ws": "   ",
                                                             "c2": "Real second-page text"}[c]}]), \
             patch.object(pdf_mod, "_detect_alignment", return_value=None):
            blocks = pdf_mod._extract_pdf_blocks_by_page(Path("dummy.pdf"))

        texts = [b["text"] for b in blocks]
        self.assertNotIn("   ", texts)  # 공백뿐인 컨테이너 자체가 블록으로 안 남아야 함
        real_second_page = next(b for b in blocks if b["text"] == "Real second-page text")
        self.assertTrue(real_second_page.get("pageBreakBefore"))


class TestCsvXlsx(Base):
    def test_roundtrip_korean(self):
        src = self.tmp / "한글데이터.csv"
        rows = [["이름", "수량"], ["김철수", "3"], ["이영희", "5"]]
        with src.open("w", newline="", encoding="utf-8-sig") as f:
            csv.writer(f).writerows(rows)
        xlsx = data.csv_to_xlsx(src, self.tmp)
        self.assertTrue(xlsx.exists())
        back_dir = self.tmp / "back"
        back_dir.mkdir()
        back = data.xlsx_to_csv(xlsx, back_dir)
        text = back.read_text(encoding="utf-8-sig")
        self.assertIn("김철수", text)
        self.assertIn("이영희", text)

    def test_xlsx_to_csv_formats_dates_and_whole_number_floats(self):
        """엑셀에서 보던 모습(날짜 "2026-07-31", 정수 "3")과 다르게 파이썬
        객체를 그대로 str()해서 "2026-07-31 00:00:00"·"3.0"처럼 나오던 문제."""
        import datetime
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.append(["날짜", "정수형태float", "일반소수", "문자열"])
        ws.append([datetime.date(2026, 7, 31), 3.0, 3.5, "그대로"])
        src = self.tmp / "d.xlsx"
        wb.save(src)

        out = data.xlsx_to_csv(src, self.tmp)
        rows = list(csv.reader(out.read_text(encoding="utf-8-sig").splitlines()))
        self.assertEqual(rows[1], ["2026-07-31", "3", "3.5", "그대로"])

    def test_xlsx_to_csv_uses_first_sheet_not_active_tab(self):
        """코드 리뷰 지적: wb.active는 "첫 번째 시트"가 아니라 파일이 마지막
        저장 시점에 열려 있던 탭이다. UI 고지 문구는 "첫 번째 시트만
        변환돼요"이므로, 두 번째 시트가 활성 탭이어도 실제로는 첫 번째
        시트가 나가야 한다."""
        from openpyxl import Workbook
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "첫번째"
        ws1.append(["첫번째시트데이터"])
        ws2 = wb.create_sheet("두번째")
        ws2.append(["두번째시트데이터"])
        wb.active = 1  # 두 번째 시트를 활성 탭으로 저장(실사용에서 흔한 상태)
        src = self.tmp / "d.xlsx"
        wb.save(src)

        out = data.xlsx_to_csv(src, self.tmp)
        text = out.read_text(encoding="utf-8-sig")
        self.assertIn("첫번째시트데이터", text)
        self.assertNotIn("두번째시트데이터", text)

    def test_xlsx_sheet_count(self):
        from openpyxl import Workbook
        wb = Workbook()
        wb.active.append(["a"])
        wb.create_sheet("두번째")
        src = self.tmp / "multi.xlsx"
        wb.save(src)
        self.assertEqual(data.xlsx_sheet_count(src), 2)

        wb2 = Workbook()
        wb2.active.append(["a"])
        src2 = self.tmp / "single.xlsx"
        wb2.save(src2)
        self.assertEqual(data.xlsx_sheet_count(src2), 1)

    def test_cp949_read(self):
        src = self.tmp / "euckr.csv"
        src.write_bytes("제품,가격\n노트북,1200000\n".encode("cp949"))
        out = data.csv_to_json(src, self.tmp)
        payload = json.loads(out.read_text(encoding="utf-8"))
        self.assertEqual(payload[0]["제품"], "노트북")


class TestCsvJson(Base):
    def test_csv_to_json_headers(self):
        src = self.tmp / "d.csv"
        src.write_text("a,b\n1,2\n3,4\n", encoding="utf-8")
        out = data.csv_to_json(src, self.tmp)
        self.assertEqual(json.loads(out.read_text()), [{"a": "1", "b": "2"}, {"a": "3", "b": "4"}])

    def test_json_to_csv_dicts_union_keys(self):
        src = self.tmp / "d.json"
        src.write_text(json.dumps([{"a": 1}, {"a": 2, "b": "x"}]), encoding="utf-8")
        out = data.json_to_csv(src, self.tmp)
        rows = list(csv.reader(out.read_text(encoding="utf-8-sig").splitlines()))
        self.assertEqual(rows[0], ["a", "b"])
        self.assertEqual(rows[2], ["2", "x"])

    def test_json_bad_shape(self):
        src = self.tmp / "d.json"
        src.write_text('{"not": "a list"}', encoding="utf-8")
        with self.assertRaises(ConversionError) as ctx:
            data.json_to_csv(src, self.tmp)
        self.assertEqual(ctx.exception.key, "err.jsonshape")

    def test_csv_to_json_preserves_embedded_newline_and_quotes(self):
        """셀 안에 줄바꿈·이스케이프된 큰따옴표가 있으면 값이 잘리던 버그
        (text.splitlines() 선분할 + Sniffer의 doublequote 오탐)."""
        src = self.tmp / "d.csv"
        with src.open("w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerows([
                ["이름", "메모"],
                ["김철수", '비고: "특이사항, 있음"'],
                ["이영희", "여러줄\n텍스트 포함"],
            ])
        out = data.csv_to_json(src, self.tmp)
        payload = json.loads(out.read_text(encoding="utf-8"))
        self.assertEqual(payload[0]["메모"], '비고: "특이사항, 있음"')
        self.assertEqual(payload[1]["메모"], "여러줄\n텍스트 포함")

    def test_csv_to_json_semicolon_delimiter(self):
        src = self.tmp / "d.csv"
        src.write_text("a;b\n1;2\n", encoding="utf-8")
        out = data.csv_to_json(src, self.tmp)
        self.assertEqual(json.loads(out.read_text()), [{"a": "1", "b": "2"}])


def _block_text(block: dict) -> str:
    """블록의 runs를 이어붙인 평문(DEC-038 — docx_to_blocks가 이제 항상
    runs를 낸다). 마커·본문이 서로 다른 run에 나뉘어 있어도 하나로 합친다."""
    return "".join(r["text"] for r in block["runs"])


class TestDocxExtractNumbering(Base):
    """코드 리뷰 지적: DOCX 자동 번호·불릿은 numbering.xml 서식일 뿐 문단
    텍스트가 아니므로, item.text만 추출하면 눈에 보이는 마커가 사라진다."""

    def test_style_based_numbered_and_bullet_list(self):
        from docx import Document
        src = self.tmp / "d.docx"
        doc = Document()
        doc.add_paragraph("일반 문단")
        doc.add_paragraph("첫 항목", style="List Number")
        doc.add_paragraph("둘째 항목", style="List Number")
        doc.add_paragraph("불릿 항목", style="List Bullet")
        doc.save(src)

        blocks = docx_to_blocks(src)
        texts = [_block_text(b) for b in blocks]
        self.assertEqual(texts, [
            "일반 문단", "1. 첫 항목", "2. 둘째 항목", "• 불릿 항목",
        ])

    def test_direct_numpr_without_named_style(self):
        """사용자가 툴바로 번호 매기기를 켠 경우 — 문단 자신의 pPr에 numPr이
        직접 있고(스타일 경유 아님), 목록 사이 일반 문단이 있어도 같은 numId의
        카운터는 이어진다."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def add_numbered(doc, text, num_id="1", ilvl="0"):
            p = doc.add_paragraph(text)
            ppr = p._p.get_or_add_pPr()
            numpr = OxmlElement("w:numPr")
            e_ilvl = OxmlElement("w:ilvl")
            e_ilvl.set(qn("w:val"), ilvl)
            e_num = OxmlElement("w:numId")
            e_num.set(qn("w:val"), num_id)
            numpr.append(e_ilvl)
            numpr.append(e_num)
            ppr.append(numpr)
            return p

        src = self.tmp / "d.docx"
        doc = Document()
        add_numbered(doc, "목록 1")
        add_numbered(doc, "목록 2")
        doc.add_paragraph("목록 사이 일반 문단")
        add_numbered(doc, "목록 3")
        doc.save(src)

        blocks = docx_to_blocks(src)
        texts = [_block_text(b) for b in blocks]
        self.assertEqual(texts[2], "목록 사이 일반 문단")
        # numId=1의 서식(decimal/bullet 등)은 기본 템플릿에 따라 달라질 수
        # 있으므로 마커 문자 자체보다 "일반 문단은 그대로, 목록 항목에는
        # (스타일 경유 없이도) 접두어가 붙는다"는 불변식을 확인한다.
        self.assertNotEqual(texts[0], "목록 1")
        self.assertTrue(texts[0].endswith("목록 1"))
        self.assertNotEqual(texts[3], "목록 3")
        self.assertTrue(texts[3].endswith("목록 3"))


class TestDocxExtractAlignment(Base):
    """DEC-040: 문단에 직접 지정된 정렬만 읽는다(스타일 상속은 범위 밖).
    명시적으로 지정 안 된 문단은 Word가 실제로 렌더링하는 값(왼쪽)을
    "left"로 명시한다 — 자동 리뷰로 발견된 회귀 수정: 예전엔 이 경우
    "align" 필드 자체를 생략했는데, HWP 쪽(JsonToHwp)이 "정렬 미지정"을
    문서 기본 ParaShape(양쪽 정렬)로 해석해 평범한 왼쪽 정렬 문단이 전부
    양쪽 정렬로 바뀌는 문제가 있었다."""

    def test_explicit_alignment_extracted(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        src = self.tmp / "d.docx"
        doc = Document()
        doc.add_paragraph("기본")
        p_center = doc.add_paragraph("가운데")
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right = doc.add_paragraph("오른쪽")
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_left = doc.add_paragraph("왼쪽 명시")
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_justify = doc.add_paragraph("양쪽")
        p_justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.save(src)

        blocks = docx_to_blocks(src)
        by_text = {b["runs"][0]["text"]: b.get("align") for b in blocks}
        self.assertEqual(by_text["기본"], "left")
        self.assertEqual(by_text["가운데"], "center")
        self.assertEqual(by_text["오른쪽"], "right")
        self.assertEqual(by_text["왼쪽 명시"], "left")
        self.assertEqual(by_text["양쪽"], "justify")


class TestPdfAlignmentClassification(unittest.TestCase):
    """DEC-040: _classify_alignment는 pdfminer 객체 없이 줄별 bbox만으로
    판정하는 순수 함수라 hand-crafted bbox로 직접 검증할 수 있다(실제
    pdfminer 파싱을 거친 end-to-end 검증은 tests/test_format_fidelity.py에
    별도로 있음)."""

    PAGE_W = 612.0

    def test_single_line_no_strong_signal_treated_as_left(self):
        self.assertEqual(_classify_alignment([(72, 700, 200, 712)], self.PAGE_W), "left")

    def test_single_line_centered(self):
        self.assertEqual(_classify_alignment([(206, 700, 406, 712)], self.PAGE_W), "center")

    def test_single_line_right_not_detected_treated_as_left(self):
        """한 줄만으로는 오른쪽 정렬을 판단하지 않는다 — "오른쪽 여백이
        작다"가 문서의 정상적인 여백인지 진짜 오른쪽 정렬인지 한 줄만
        봐서는 구분할 근거가 없다(로컬 검증 중 발견해 범위를 좁힘). 오른쪽
        정렬로 확정하지 않는 대신 "left"를 명시적으로 돌려준다 — None을
        반환하면 HWP 쪽에서 문서 기본 정렬(양쪽 정렬)로 해석되는 회귀가
        있었다(자동 리뷰로 발견)."""
        self.assertEqual(_classify_alignment([(400, 700, 540, 712)], self.PAGE_W), "left")

    def test_single_line_flush_with_left_edge_returns_none(self):
        self.assertIsNone(_classify_alignment([(0, 700, 200, 712)], self.PAGE_W))

    def test_multi_line_left_aligned_treated_as_left(self):
        boxes = [(72, 700, 300, 712), (72, 680, 500, 692), (72, 660, 150, 672)]
        self.assertEqual(_classify_alignment(boxes, self.PAGE_W), "left")

    def test_two_line_left_aligned_not_misclassified_as_justify(self):
        """자동 리뷰로 발견된 회귀: 마지막 줄을 뺀 "본문 줄"이 1개뿐인 2줄
        문단은 그 한 줄의 오른쪽 끝을 자기 자신과 비교해 항상 "일치"로
        판정돼, 평범한 2줄 왼쪽 정렬 문단이 양쪽 정렬로 잘못 분류됐다."""
        boxes = [(72, 700, 300, 712), (72, 680, 500, 692)]
        self.assertEqual(_classify_alignment(boxes, self.PAGE_W), "left")

    def test_two_line_right_aligned_not_detected(self):
        """본문 줄이 1개뿐이면 오른쪽·양쪽 정렬 판정 자체를 포기한다(위
        회귀 수정의 트레이드오프) — 2줄 오른쪽 정렬은 이제 감지되지 않고
        None(판단 근거 부족)을 돌려준다. 잘못된 확정보다 안전하다는 이
        함수의 기본 원칙과 일치."""
        boxes = [(300, 700, 540, 712), (100, 680, 540, 692)]
        self.assertIsNone(_classify_alignment(boxes, self.PAGE_W))

    def test_multi_line_right_aligned(self):
        boxes = [(300, 700, 540, 712), (100, 680, 540, 692), (400, 660, 540, 672)]
        self.assertEqual(_classify_alignment(boxes, self.PAGE_W), "right")

    def test_multi_line_justified_last_line_short(self):
        boxes = [(72, 700, 540, 712), (72, 680, 540, 692), (72, 660, 200, 672)]
        self.assertEqual(_classify_alignment(boxes, self.PAGE_W), "justify")

    def test_multi_line_centered(self):
        boxes = [(200, 700, 412, 712), (150, 680, 462, 692), (250, 660, 362, 672)]
        self.assertEqual(_classify_alignment(boxes, self.PAGE_W), "center")

    def test_multi_line_ambiguous_returns_none(self):
        boxes = [(50, 700, 200, 712), (300, 680, 500, 692), (80, 660, 550, 672)]
        self.assertIsNone(_classify_alignment(boxes, self.PAGE_W))

    def test_empty_boxes_returns_none(self):
        self.assertIsNone(_classify_alignment([], self.PAGE_W))


class TestDocxExtractCharFormatting(Base):
    """DEC-038: docx_extract.py가 run별 문자 서식(굵게/기울임/밑줄/크기/색상)을
    추출하는지 — JDK 없이 순수 파이썬 추출 로직만 검증(HWP 왕복까지 포함한
    통합 테스트는 test_pipeline.py::TestHwp에 있음, hwplib 필요)."""

    def test_bold_italic_underline_size_color_extracted_per_run(self):
        from docx import Document
        from docx.shared import Pt, RGBColor

        src = self.tmp / "d.docx"
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("일반 ")
        bold_run = p.add_run("굵게")
        bold_run.bold = True
        styled_run = p.add_run("기울임크게빨강")
        styled_run.italic = True
        styled_run.font.size = Pt(18)
        styled_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        doc.save(src)

        blocks = docx_to_blocks(src)
        self.assertEqual(len(blocks), 1)
        runs = blocks[0]["runs"]
        self.assertEqual(runs[0], {"text": "일반 ", "bold": False, "italic": False,
                                    "underline": False, "size": None, "color": None})
        self.assertEqual(runs[1], {"text": "굵게", "bold": True, "italic": False,
                                    "underline": False, "size": None, "color": None})
        self.assertEqual(runs[2], {"text": "기울임크게빨강", "bold": False, "italic": True,
                                    "underline": False, "size": 18.0, "color": "FF0000"})

    def test_plain_run_has_no_formatting(self):
        from docx import Document
        src = self.tmp / "d.docx"
        doc = Document()
        doc.add_paragraph("서식 없는 문단")
        doc.save(src)

        blocks = docx_to_blocks(src)
        self.assertEqual(blocks[0]["runs"], [
            {"text": "서식 없는 문단", "bold": False, "italic": False,
             "underline": False, "size": None, "color": None},
        ])

    def test_numbering_marker_prepended_as_unformatted_run(self):
        """마커(번호/불릿)는 본문 run의 서식과 무관하게 별도의 서식 없는
        run으로 앞에 붙는다 — 본문이 굵게라도 마커까지 굵어지지 않는다."""
        from docx import Document

        src = self.tmp / "d.docx"
        doc = Document()
        p = doc.add_paragraph("", style="List Number")
        bold_run = p.add_run("굵은 항목")
        bold_run.bold = True
        doc.save(src)

        blocks = docx_to_blocks(src)
        runs = blocks[0]["runs"]
        self.assertEqual(runs[0]["text"], "1. ")
        self.assertFalse(runs[0]["bold"])
        self.assertEqual(runs[1], {"text": "굵은 항목", "bold": True, "italic": False,
                                    "underline": False, "size": None, "color": None})

    def test_hyperlink_only_paragraph_not_dropped(self):
        """`paragraph.runs`(python-docx)는 <w:hyperlink> 안에 중첩된 run을
        포함하지 않는다 — 문단 전체가 하이퍼링크 하나뿐이면 옛 구현은 runs가
        빈 리스트가 되어 문단째로 조용히 드롭됐다(회귀, 이번에 수정)."""
        import docx.oxml
        import docx.opc.constants
        from docx import Document
        from docx.oxml.ns import qn

        src = self.tmp / "d.docx"
        doc = Document()
        p = doc.add_paragraph()
        r_id = p.part.relate_to(
            "https://example.com",
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        hyperlink = docx.oxml.OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run_el = docx.oxml.OxmlElement("w:r")
        t = docx.oxml.OxmlElement("w:t")
        t.text = "링크텍스트"
        run_el.append(t)
        hyperlink.append(run_el)
        p._p.append(hyperlink)
        doc.save(src)

        blocks = docx_to_blocks(src)
        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0]["runs"][0]["text"], "링크텍스트")


class TestDocxBuildFont(Base):
    """DEC-015: 생성 DOCX는 한글 글꼴을 모든 run에 명시해야 한다 — 실사용 중
    글자 깨짐이 재현된 근본 원인(글꼴 미지정 → 뷰어별 대체 글꼴 불일치)."""

    def test_paragraph_and_table_runs_declare_east_asian_font(self):
        blocks = [
            {"type": "p", "text": "한글 문단"},
            {"type": "table", "rows": [["셀1", "셀2"]]},
        ]
        out = blocks_to_docx(blocks, self.tmp / "out.docx")

        import zipfile
        with zipfile.ZipFile(out) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        self.assertIn(f'w:eastAsia="{EAST_ASIAN_FONT}"', xml)
        self.assertIn(f'w:ascii="{EAST_ASIAN_FONT}"', xml)

        from docx import Document
        doc = Document(out)
        for p in doc.paragraphs:
            for run in p.runs:
                self.assertEqual(run.font.name, EAST_ASIAN_FONT)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            self.assertEqual(run.font.name, EAST_ASIAN_FONT)


class TestDocxBuildAlignment(Base):
    """DEC-040: "align"이 있는 블록만 명시적으로 정렬을 설정하고, 없으면
    DOCX 기본 정렬(왼쪽, python-docx의 None)을 그대로 둔다."""

    def test_align_field_sets_paragraph_alignment(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        blocks = [
            {"type": "p", "text": "기본"},
            {"type": "p", "text": "가운데", "align": "center"},
            {"type": "p", "text": "오른쪽", "align": "right"},
            {"type": "p", "text": "왼쪽", "align": "left"},
            {"type": "p", "text": "양쪽", "align": "justify"},
        ]
        out = blocks_to_docx(blocks, self.tmp / "out.docx")
        doc = Document(out)
        by_text = {p.text: p.alignment for p in doc.paragraphs}
        self.assertIsNone(by_text["기본"])
        self.assertEqual(by_text["가운데"], WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(by_text["오른쪽"], WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(by_text["왼쪽"], WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(by_text["양쪽"], WD_ALIGN_PARAGRAPH.JUSTIFY)


class TestOutputNaming(Base):
    def test_auto_rename_never_overwrites(self):
        (self.tmp / "r.pdf").write_text("existing")
        p1 = unique_output_path(self.tmp, "r", "pdf")
        self.assertEqual(p1.name, "r (1).pdf")
        p1.write_text("second")
        p2 = unique_output_path(self.tmp, "r", "pdf")
        self.assertEqual(p2.name, "r (2).pdf")

    def test_finalize_concurrent_same_stem_never_overwrites(self):
        """숨은 버그(내부 감사): 같은 파일을 두 번 추가해 같은 포맷으로 동시
        변환하면(QThreadPool, 최대 4개 동시 — workers.py) 두 워커가 거의
        동시에 finalize()를 호출한다. "이름 충돌 확인 → 이동"이 원자적이지
        않으면(TOCTOU) 둘 다 같은 "충돌 없음" 경로를 계산해 한쪽이 다른 쪽
        결과물을 조용히 덮어쓴다 — 실제 스레드로 재현 확인 후 락으로 수정."""
        import threading
        from app.output import finalize

        source = self.tmp / "r.docx"
        source.write_text("원본")
        tmp_root = Path(tempfile.mkdtemp())
        self.addCleanup(shutil.rmtree, tmp_root, ignore_errors=True)

        results = []
        barrier = threading.Barrier(4)

        def worker(label):
            tmp_out = tmp_root / f"produced_{label}.pdf"
            tmp_out.write_text(label)
            barrier.wait()
            out, renamed = finalize(tmp_out, source, "pdf")
            results.append(str(out))

        threads = [threading.Thread(target=worker, args=(c,)) for c in "ABCD"]
        for t in threads:
            t.start()
        for t in threads:
            t.join()

        self.assertEqual(len(set(results)), 4)  # 네 결과 경로가 전부 달라야 함
        pdfs = [p for p in self.tmp.iterdir() if p.suffix == ".pdf"]
        self.assertEqual(len(pdfs), 4)  # 실제로 4개 파일 모두 생성(덮어쓰기 없음)


if __name__ == "__main__":
    unittest.main()
