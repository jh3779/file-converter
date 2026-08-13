"""PDF→DOCX 이미지·벡터 도형(사각형/직선) 추출 테스트 — 이미지·표 테두리
반영 개선.

test_pdf_to_docx_layout.py는 텍스트 레이어만 다루는 손으로 만든 최소 PDF를
쓴다(당시엔 이미지·표 테두리가 DEC-037 범위 밖이었음). 이 파일은
test_pdf_to_pptx_visuals.py와 똑같은 "손으로 최소 PDF 바이트를 직접 조립"
방식(content-stream 연산자 re/m·l/Do까지 확장)으로, LTRect(사각형 채움)·
LTLine(직선)·LTImage(이미지 XObject)가 DOCX에서도 원래 위치·크기·색상으로
재구성되는지 검증한다 — python-docx엔 pptx 같은 고수준 도형 API가 없어
문단 테두리(w:pBdr)/채움(w:shd)으로 표현하므로, 검증도 원시 XML을 직접
파싱해서 한다.
"""
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app import converters

EMU_PER_PT = 12700
TWIPS_PER_PT = 20


def _build_pdf(path: Path, content_ops: str, extra_resources: str = "",
                extra_objects: list[bytes] | None = None, width=612, height=792):
    """단일 페이지 최소 PDF. test_pdf_to_pptx_visuals.py의 동명 헬퍼와 동일."""
    extra_objects = extra_objects or []
    content = content_ops.encode()
    stream = b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream"

    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {width} {height}] /Contents 4 0 R"
         f" /Resources << {extra_resources} >> >>").encode(),
        stream,
        *extra_objects,
    ]

    buf = b"%PDF-1.4\n"
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref = len(buf)
    buf += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        buf += f"{off:010d} 00000 n \n".encode()
    buf += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    path.write_bytes(buf)


def _frame_pos(paragraph):
    """문단의 w:framePr에서 (x_pt, y_pt, w_pt, h_pt)를 뽑는다."""
    pPr = paragraph._p.find(qn("w:pPr"))
    frame = pPr.find(qn("w:framePr"))
    def val(attr):
        return int(frame.get(qn(f"w:{attr}"))) / TWIPS_PER_PT
    return val("x"), val("y"), val("w"), val("h")


class TestPdfToDocxVisuals(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_filled_rect_extracted_with_position_and_shading(self):
        """0 0 1 rg(파랑 채움) 100 600 200 50 re f — 채워진 사각형 하나."""
        src = self.tmp / "rect.pdf"
        _build_pdf(src, "0 0 1 rg 100 600 200 50 re f")
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        shaded = [p for p in doc.paragraphs
                  if p._p.find(qn("w:pPr")) is not None
                  and p._p.find(qn("w:pPr")).find(qn("w:shd")) is not None]
        self.assertEqual(len(shaded), 1)
        p = shaded[0]
        shd = p._p.find(qn("w:pPr")).find(qn("w:shd"))
        self.assertEqual(shd.get(qn("w:fill")), "0000FF")

        x, y, w, h = _frame_pos(p)
        self.assertAlmostEqual(x, 100, delta=1)
        self.assertAlmostEqual(w, 200, delta=1)
        self.assertAlmostEqual(h, 50, delta=1)

    def test_stroked_line_extracted_with_position_and_color(self):
        """1 0 0 RG(빨강 선) 100 600 m 300 600 l S — 수평 직선 하나."""
        src = self.tmp / "line.pdf"
        _build_pdf(src, "1 0 0 RG 2 w 100 600 m 300 600 l S")
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        bordered = [p for p in doc.paragraphs
                    if p._p.find(qn("w:pPr")) is not None
                    and p._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None]
        self.assertEqual(len(bordered), 1)
        p = bordered[0]
        pBdr = p._p.find(qn("w:pPr")).find(qn("w:pBdr"))
        # 수평선이므로 bottom 변 하나에만 테두리가 있어야 한다(이중선으로
        # 보이지 않도록 변 하나만 쓰는 설계, 실제 렌더링으로 확인).
        self.assertIsNotNone(pBdr.find(qn("w:bottom")))
        self.assertIsNone(pBdr.find(qn("w:top")))
        self.assertIsNone(pBdr.find(qn("w:left")))
        self.assertEqual(pBdr.find(qn("w:bottom")).get(qn("w:color")), "FF0000")

        x, y, w, h = _frame_pos(p)
        self.assertAlmostEqual(x, 100, delta=1)
        self.assertAlmostEqual(w, 200, delta=1)

    def test_image_xobject_extracted_and_embedded(self):
        """가공 없는 2x2 RGB 원시 이미지 XObject를 /Im1로 참조해 (50,600)
        위치에 100x100pt 크기로 배치 — 결과 DOCX에 그림 run으로 나와야 한다
        (픽셀 값 자체는 검증하지 않음, 위치·존재만 확인)."""
        src = self.tmp / "image.pdf"
        raw_rgb = bytes([
            255, 0, 0,    0, 255, 0,
            0, 0, 255,    255, 255, 0,
        ])
        image_obj = (
            b"<< /Type /XObject /Subtype /Image /Width 2 /Height 2 "
            b"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Length " +
            str(len(raw_rgb)).encode() + b" >>\nstream\n" + raw_rgb + b"\nendstream"
        )
        _build_pdf(
            src,
            "q 100 0 0 100 50 600 cm /Im1 Do Q",
            extra_resources="/XObject << /Im1 5 0 R >>",
            extra_objects=[image_obj],
        )
        out = converters.convert(src, "docx", self.tmp)

        doc = Document(out)
        pic_paragraphs = [p for p in doc.paragraphs
                          if p._p.findall(".//" + qn("pic:pic"))]
        self.assertEqual(len(pic_paragraphs), 1)
        x, y, w, h = _frame_pos(pic_paragraphs[0])
        self.assertAlmostEqual(x, 50, delta=1)
        self.assertAlmostEqual(w, 100, delta=1)
        self.assertAlmostEqual(h, 100, delta=1)

    def test_text_still_renders_alongside_visuals(self):
        """도형·텍스트가 섞인 페이지에서 텍스트 줄도 그대로 나오는지(회귀
        방지 — visuals 처리가 기존 텍스트 처리를 깨지 않아야 한다)."""
        src = self.tmp / "mixed.pdf"
        _build_pdf(src, "0 0 1 rg 100 700 200 20 re f BT /F1 12 Tf 100 600 Td (hello) Tj ET",
                   extra_resources="/Font << /F1 6 0 R >>",
                   extra_objects=[b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"])
        out = converters.convert(src, "docx", self.tmp)
        doc = Document(out)
        # 손으로 만든 최소 PDF엔 폰트 너비(Widths) 정보가 없어 pdfminer가
        # 글자를 하나씩 별도 줄로 나누는 경우가 있다 — 여기서는 "글자가
        # 실제로 옮겨졌는지"만 확인한다(줄 단위 세부 배치는 다른 테스트가 검증).
        joined = "".join(p.text for p in doc.paragraphs)
        self.assertEqual(joined, "hello")


if __name__ == "__main__":
    unittest.main()
