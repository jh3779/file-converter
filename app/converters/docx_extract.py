"""DOCX → 구조 블록 (python-docx). DOCX→HWP 파이프라인 1단계 (docx_build.blocks_to_docx의 역방향).

블록 형식: {"type":"p","runs":[{"text":str,"bold":bool,"italic":bool,
"underline":bool,"size":float,"color":"RRGGBB"}, ...],"align":"left"|"center"|"right"|"justify"(선택)} |
{"type":"table","rows":[[{"text":str,"colSpan":int,"rowSpan":int},...],...],"colWidthsMm":[float,...]}
문서 순서(본문에 등장하는 순서)대로 문단·표를 함께 추출한다 — python-docx는
document.paragraphs/document.tables를 각각 따로 주기 때문에 body XML을 직접
순회해야 순서가 보존된다.

문자 서식(굵게/기울임/밑줄/크기/색상, DEC-038)은 run 단위로 그대로 추출한다
— run.bold/italic/underline이 None이면(스타일에서 상속받고 직접 지정은 없는
경우) 미지정으로 보고 False로 안전하게 처리한다. HwpToJson.java(HWP→DOCX
읽기 방향, DEC-027)와 대칭이지만 스타일 상속 체인 전체를 해석하지는
않는다는 점은 같은 단순화 원칙(번호 매기기의 restart 규칙 미재현과 동일).
표 셀 텍스트는 이번 범위 밖이라 문자 서식 없이 평문으로 남는다.

문단 정렬(DEC-040): 문단에 직접 지정된 정렬(w:jc)만 읽는다(스타일 상속은
범위 밖). 직접 지정이 없으면 Word가 실제로 렌더링하는 값인 "left"를
명시한다 — 예전엔 이 경우 "align" 필드 자체를 생략했는데, HWP 쪽
(JsonToHwp.java)이 "정렬 미지정"을 문서 기본 ParaShape(양쪽 정렬)로
해석해 정렬을 전혀 지정하지 않은 평범한 DOCX 문단이 전부 양쪽 정렬로
바뀌는 회귀가 있었다(자동 리뷰로 발견, 실제 hwplib DOCX→HWP→JSON
왕복으로 재현 확인 후 수정).

표 셀 병합·열 너비(DEC-035): python-docx의 `table.cell(r, c)`는 병합된
영역의 모든 그리드 위치에서 **같은 `_tc`(내부 XML 요소) 객체**를 돌려준다
(가로 병합·세로 병합 둘 다, 직접 만든 병합 DOCX로 실측 확인) — 이 객체
동일성으로 병합 영역을 감지한다. 각 행의 `rows[r]`는 그 행에서 처음
등장하는(=병합의 왼쪽 위 모서리인) 셀만 담고, 이전 행의 세로 병합이
차지한 칸은 건너뛴다 — 소비하는 쪽(JsonToHwp.java)이 HTML 표 렌더링과
같은 원리로 "위에서 내려오는 rowSpan이 점유한 칸"을 추적하며 채워야
한다. 열 너비는 `table.columns[i].width`(EMU)를 mm로 변환해 그대로
전달 — DOCX table.add_table()로 만든 문서도 이 값이 항상 채워져 있음을
확인(명시적으로 지정 안 하면 기본값이 채워짐, None이 되는 경우 없음).

번호·불릿 목록 주의: DOCX의 자동 번호("1.", "가.")·불릿("•")은 문단의 실제
텍스트(w:t)가 아니라 numbering.xml에 정의된 서식이 뷰어에 의해 화면에만
그려지는 것이라, item.text만 추출하면 눈에 보이는 마커가 조용히 사라진다
(코드 리뷰 지적, 실제 재현 확인 후 보완). 소수(decimal)·로마자·알파벳·불릿
서식은 numbering.xml을 해석해 마커 문자열을 만들어 문단 맨 앞에 서식 없는
run으로 붙인다(마커 자체는 numbering.xml의 별도 서식을 갖지만 이번 범위
에서는 재현하지 않음 — 문서화된 단순화). 다단계 중첩 목록의 상위 레벨
변경 시 하위 레벨 카운터 재시작 등 OOXML 번호 매기기의 전체 규칙(요소
restart·startOverride 등)까지는 재현하지 않는다 — (numId, ilvl) 쌍별로
문서 순서대로 단순 증가하는 카운터를 쓴다. 지원하지 않는 서식(예: 사용자
정의 다단계 서식)은 마커 없이 본문만 유지한다(내용 유실은 없음 — 마커만 생략).
"""
from pathlib import Path

_ROMAN_VALUES = (
    (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
    (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
    (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
)


def _to_roman(n: int) -> str:
    out = []
    for value, symbol in _ROMAN_VALUES:
        count, n = divmod(n, value)
        out.append(symbol * count)
    return "".join(out)


def _to_letter(n: int) -> str:
    letters = []
    while n > 0:
        n, rem = divmod(n - 1, 26)
        letters.append(chr(ord("a") + rem))
    return "".join(reversed(letters))


def _format_marker(numfmt: str, lvltext: str, n: int) -> str | None:
    if numfmt == "bullet":
        return "•"
    if numfmt == "decimal":
        value = str(n)
    elif numfmt == "decimalZero":
        value = f"{n:02d}"
    elif numfmt == "upperRoman":
        value = _to_roman(n)
    elif numfmt == "lowerRoman":
        value = _to_roman(n).lower()
    elif numfmt == "upperLetter":
        value = _to_letter(n).upper()
    elif numfmt == "lowerLetter":
        value = _to_letter(n)
    else:
        return None  # 지원하지 않는 서식 — 마커 생략(본문 텍스트는 유지)
    if lvltext and "%" in lvltext:
        import re
        return re.sub(r"%\d+", value, lvltext, count=1)
    return value + "."


def _numbering_levels(document):
    """{numId: {ilvl: (numFmt, lvlText, start)}} — numId는 문자열(w:val 원문)."""
    from docx.oxml.ns import qn

    try:
        numbering_el = document.part.numbering_part.element
    except Exception:
        return {}

    abstract_levels: dict[str, dict[str, tuple]] = {}
    for abstract_num in numbering_el.findall(qn("w:abstractNum")):
        abstract_id = abstract_num.get(qn("w:abstractNumId"))
        levels = {}
        for lvl in abstract_num.findall(qn("w:lvl")):
            ilvl = lvl.get(qn("w:ilvl"))
            numfmt_el = lvl.find(qn("w:numFmt"))
            lvltext_el = lvl.find(qn("w:lvlText"))
            start_el = lvl.find(qn("w:start"))
            numfmt = numfmt_el.get(qn("w:val")) if numfmt_el is not None else "decimal"
            lvltext = lvltext_el.get(qn("w:val")) if lvltext_el is not None else "%1."
            try:
                start = int(start_el.get(qn("w:val"))) if start_el is not None else 1
            except (TypeError, ValueError):
                start = 1
            levels[ilvl] = (numfmt, lvltext, start)
        abstract_levels[abstract_id] = levels

    result: dict[str, dict[str, tuple]] = {}
    for num in numbering_el.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abstract_id_el = num.find(qn("w:abstractNumId"))
        if abstract_id_el is None:
            continue
        abstract_id = abstract_id_el.get(qn("w:val"))
        if abstract_id in abstract_levels:
            result[num_id] = abstract_levels[abstract_id]
    return result


def _paragraph_numpr(paragraph):
    """문단 자신 또는(직접 지정이 없으면) 적용된 스타일 체인에서 numPr을 찾는다.
    "List Number"·"List Bullet" 같은 내장 스타일은 문단이 아니라 스타일 쪽에
    numPr이 있다 — 실제 생성 DOCX로 확인함."""
    from docx.oxml.ns import qn

    ppr = paragraph._p.find(qn("w:pPr"))
    numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
    if numpr is None:
        style = paragraph.style
        seen = set()
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            style_ppr = style.element.find(qn("w:pPr"))
            if style_ppr is not None:
                numpr = style_ppr.find(qn("w:numPr"))
                if numpr is not None:
                    break
            style = getattr(style, "base_style", None)
    if numpr is None:
        return None
    num_id_el = numpr.find(qn("w:numId"))
    ilvl_el = numpr.find(qn("w:ilvl"))
    if num_id_el is None:
        return None
    num_id = num_id_el.get(qn("w:val"))
    if num_id == "0":  # 0은 "목록 아님"(번호 제거)을 뜻하는 예약값
        return None
    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
    return num_id, ilvl


_ALIGN_TO_STR = {
    "LEFT": "left", "CENTER": "center", "RIGHT": "right",
    "JUSTIFY": "justify", "JUSTIFY_MED": "justify", "JUSTIFY_HI": "justify",
    "JUSTIFY_LOW": "justify", "DISTRIBUTE": "justify", "THAI_JUSTIFY": "justify",
}


def _paragraph_align(paragraph) -> str | None:
    """문단에 직접 지정된 정렬만 읽는다(스타일에서 물려받는 값은 안 봄 —
    style.paragraph_format.alignment까지 걸어야 하는데, 이 프로젝트가
    다루는 실사용 문서에서 정렬은 거의 항상 문단에 직접 지정돼 있어
    범위 밖으로 둠, DEC-040). alignment가 None이면 Word가 실제로 렌더링하는
    값(왼쪽)을 명시적으로 "left"로 돌려준다 — 예전엔 이 필드 자체를
    생략했었는데, 그러면 HWP 쪽(JsonToHwp)이 "정렬 미지정"을 문서 기본
    ParaShape(양쪽 정렬)로 해석해 평범한 왼쪽 정렬 문단이 전부 양쪽 정렬로
    바뀌는 회귀가 있었다(자동 리뷰로 발견, 실제 hwplib DOCX→HWP→JSON
    왕복으로 재현 확인 후 수정)."""
    alignment = paragraph.alignment
    if alignment is None:
        return "left"
    return _ALIGN_TO_STR.get(alignment.name)


def _run_to_dict(run) -> dict | None:
    if not run.text:
        return None
    size = run.font.size
    color = None
    try:
        rgb = run.font.color.rgb if run.font.color is not None else None
        if rgb is not None:
            color = str(rgb)
    except Exception:
        pass  # 테마 색상 등 RGB로 안 떨어지는 경우 — 색상 미지정으로 안전 처리
    return {
        "text": run.text,
        "bold": bool(run.bold),
        "italic": bool(run.italic),
        "underline": bool(run.underline),
        "size": (size.pt if size is not None else None),
        "color": color,
    }


def _paragraph_runs(paragraph) -> list[dict]:
    """문단의 run별 문자 서식을 추출한다(DEC-038). 빈 텍스트 run은 건너뛴다
    (예: 필드 코드 전용 run 등 — 내용 유실 아님, 서식 붙일 텍스트가 없을 뿐).

    `paragraph.runs`는 `<w:hyperlink>` 안에 중첩된 run을 포함하지 않는다
    (python-docx가 `w:p`의 직계 자식 `w:r`만 봄) — 문단 전체가 하이퍼링크
    하나로만 이루어지면(이메일/웹 링크 등) runs가 빈 리스트가 되어 문단
    전체가 조용히 드롭되는 회귀가 있었다. `iter_inner_content()`로 Run과
    Hyperlink를 문서 순서대로 순회해 하이퍼링크 내부 run도 펼친다."""
    from docx.text.hyperlink import Hyperlink
    from docx.text.run import Run

    runs = []
    for item in paragraph.iter_inner_content():
        inner_runs = item.runs if isinstance(item, Hyperlink) else [item] if isinstance(item, Run) else []
        for run in inner_runs:
            d = _run_to_dict(run)
            if d is not None:
                runs.append(d)
    return runs


_EMU_PER_MM = 36000


def _column_widths_mm(table) -> list[float] | None:
    widths = []
    for col in table.columns:
        w = col.width
        if w is None:
            return None
        widths.append(round(w / _EMU_PER_MM, 2))
    return widths


def _table_rows_with_spans(table) -> list[list[dict]]:
    """행마다 그 행에서 처음 등장하는 셀만 담되(병합의 왼쪽 위 모서리),
    colSpan/rowSpan을 함께 기록한다. table.cell(r, c)는 그리드의 모든
    칸에 대해 값을 돌려주므로(병합 영역은 같은 _tc 객체가 반복됨) 항상
    n_cols개씩 순회할 수 있다 — 병합으로 "칸이 비는" 현상은 아예 없고,
    같은 _tc가 반복되는 것만 걸러내면 된다."""
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    row_tcs = [list(table.rows[r].cells) for r in range(n_rows)]
    seen = set()
    rows_out = []
    for r in range(n_rows):
        row_out = []
        cells = row_tcs[r]
        for c in range(min(n_cols, len(cells))):
            tc_id = id(cells[c]._tc)
            if tc_id in seen:
                continue
            seen.add(tc_id)
            col_span = 1
            while c + col_span < len(cells) and id(cells[c + col_span]._tc) == tc_id:
                col_span += 1
            row_span = 1
            while r + row_span < n_rows:
                below = row_tcs[r + row_span]
                if c < len(below) and id(below[c]._tc) == tc_id:
                    row_span += 1
                else:
                    break
            text = cells[c].text.strip().replace("\n", " ")
            entry = {"text": text, "colSpan": col_span, "rowSpan": row_span}
            row_out.append(entry)
        rows_out.append(row_out)
    return rows_out


def _iter_block_items(document):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def docx_to_blocks(src: Path) -> list[dict]:
    from docx import Document
    from docx.table import Table

    doc = Document(src)
    levels = _numbering_levels(doc)
    counters: dict[tuple, int] = {}
    blocks: list[dict] = []
    for item in _iter_block_items(doc):
        if isinstance(item, Table):
            rows = _table_rows_with_spans(item)
            if any(cell["text"] for row in rows for cell in row):
                block = {"type": "table", "rows": rows}
                col_widths = _column_widths_mm(item)
                if col_widths is not None:
                    block["colWidthsMm"] = col_widths
                blocks.append(block)
        else:
            if not item.text.strip():
                continue
            runs = _paragraph_runs(item)
            numpr = _paragraph_numpr(item)
            if numpr is not None:
                num_id, ilvl = numpr
                level = levels.get(num_id, {}).get(ilvl)
                if level is not None:
                    numfmt, lvltext, start = level
                    key = (num_id, ilvl)
                    counters[key] = counters.get(key, start - 1) + 1
                    marker = _format_marker(numfmt, lvltext, counters[key])
                    if marker:
                        runs.insert(0, {"text": f"{marker} ", "bold": False, "italic": False,
                                        "underline": False, "size": None, "color": None})
            if runs:
                block = {"type": "p", "runs": runs}
                align = _paragraph_align(item)
                if align:
                    block["align"] = align
                blocks.append(block)
    return blocks
