"""구조 블록 → DOCX 생성 (python-docx). HWP/PDF → DOCX 파이프라인 공용 (DEC-007·DEC-027).

"size":float,"color":"RRGGBB"}, ...],"align":"left"|"center"|"right"|"justify"(선택)} |
{"type":"p","text":str}(구버전 호환, 서식 없는 단일 run으로 취급) |
{"type":"table","rows":[[cell,...],...]} — cell은 {"runs":[...],"colSpan":int,
"rowSpan":int}(DEC-051, 문단 블록과 같은 runs 표현으로 셀 안 문자 서식까지
반영) 또는 구버전 호환 형태(평문 문자열, {"text":str,"colSpan":int,"rowSpan":
int} — HwpToJson.java DEC-035 시절 출력)이며, 한 행에는 그 행에서 "처음
등장하는" 셀만 담긴다(세로 병합이 위에서 내려와 차지한 칸은 생략) —
JsonToHwp.java·docx_extract.py와 같은 표현. 레이아웃(정확한 위치·다단·이미지
등)은 여전히 단순화된다 — 기대치 고지는 OQ-003/DEC-010 문안으로 UI에서 안내.
문단 문자 서식(굵게/기울임/밑줄/크기/색상)은 DEC-027부터, 표 셀 안 문자
서식은 DEC-051부터 반영된다. 문단 정렬(DEC-040)은 "align"이 있을 때만
명시적으로 설정하고, 없으면 DOCX 기본 정렬(왼쪽)을 그대로 둔다. 셀 병합은
DEC-035부터 python-docx의 cell.merge()로 재현한다(열 너비까지는 이 방향에서
아직 전달되지 않음 — HwpToJson.java가 colWidthsMm을 내지 않음, 표 전체
폭은 python-docx/Word 기본 렌더링에 맡긴다).

한글 글꼴을 모든 run에 명시적으로 지정한다(DEC-015) — python-docx 기본
스타일(Calibri)은 한글 글리프가 없고, 지정을 생략하면 뷰어·OS별로 대체
글꼴이 달라져 렌더링이 일관되지 않는다(실사용 중 글자 깨짐 재현·확인됨).

글꼴은 번들 "Noto Sans KR"(OFL-1.1, engine/libreoffice에 동봉)을 쓴다.
**이 보장 범위는 우리 앱 자신의 렌더링 경로(DOCX/HWP→PDF)에 한정된다** —
그 경로에서는 이 폰트가 항상 사용 가능함을 실제로 검증했다. 결과 DOCX를
사용자가 자신의 Word/한글(HWP)에서 직접 열 때는 그 프로그램에 "Noto Sans
KR"이 설치돼 있지 않으면 다른 대체가 일어난다 — 대안으로 어느 Windows에나
있는 "맑은 고딕"을 1순위로 지정하는 시도를 로컬에서 검증했으나, 그 폰트가
없는 환경(예: 개발용 mac)에서 LibreOffice의 대체 로직이 무지정 상태보다
더 나쁜 대체를 골라 희귀 글자(뷁 등)의 매핑이 깨지는 회귀가 실측 확인되어
채택하지 않았다. 검증되지 않은 "개선"보다 검증된 현재 상태를 유지한다
(docs/06_open_questions.md 리스크 표에 외부 뷰어 잔여 리스크로 기록).
"""
from pathlib import Path

EAST_ASIAN_FONT = "Noto Sans KR"

_ALIGN_FROM_STR = None  # 지연 초기화 — docx.enum.text 임포트 비용을 문서 안 열 때는 안 지불하도록


def _align_map():
    global _ALIGN_FROM_STR
    if _ALIGN_FROM_STR is None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        _ALIGN_FROM_STR = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
    return _ALIGN_FROM_STR


def _set_font(run):
    run.font.name = EAST_ASIAN_FONT
    # python-docx는 run.font.name을 서양(ascii) 글꼴에만 반영한다 — 동아시아
    # 문자 렌더링에 쓰이는 w:eastAsia는 별도로 XML에 직접 지정해야 한다.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    if rfonts is None:
        from docx.oxml.ns import qn
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    from docx.oxml.ns import qn
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), EAST_ASIAN_FONT)


def _apply_run_style(run, run_dict: dict):
    """서식 필드(굵게/기울임/밑줄/크기/색상)를 run에 반영한다 (DEC-027).
    값이 없으면(구버전 호환 run) 아무 것도 건드리지 않고 기본값을 따른다."""
    if run_dict.get("bold") is not None:
        run.font.bold = bool(run_dict["bold"])
    if run_dict.get("italic") is not None:
        run.font.italic = bool(run_dict["italic"])
    if run_dict.get("underline") is not None:
        run.font.underline = bool(run_dict["underline"])
    size = run_dict.get("size")
    if size:
        from docx.shared import Pt
        run.font.size = Pt(size)
    color = run_dict.get("color")
    if color and color.upper() != "000000":  # 검정은 DOCX 기본값이라 굳이 명시 안 함
        from docx.shared import RGBColor
        try:
            run.font.color.rgb = RGBColor.from_string(color)
        except ValueError:
            pass  # 색상 문자열이 6자리 hex가 아니면 조용히 건너뜀(텍스트 보존 우선)


def _runs_for(block: dict) -> list[dict]:
    """블록에서 run 목록을 얻는다 — 신버전("runs")·구버전("text") 스키마 모두 지원."""
    runs = block.get("runs")
    if runs is not None:
        return runs
    text = block.get("text")
    return [{"text": text}] if text else []


def _cell_runs(cell) -> list[dict]:
    """셀 하나에서 run 목록을 얻는다(DEC-051) — 신버전({"runs":[...]})과
    구버전(평문 문자열, {"text":str}) 모두 지원. 문단 블록의 _runs_for()와
    같은 원칙."""
    if isinstance(cell, str):
        return [{"text": cell}] if cell else []
    runs = cell.get("runs")
    if runs is not None:
        return runs
    text = cell.get("text")
    return [{"text": text}] if text else []


def _place_cells_with_spans(rows: list) -> tuple[list[dict], int, int]:
    """행마다 "그 행에서 처음 등장하는 셀"만 담긴 rows(HwpToJson.java와 같은
    표현, 평문 문자열·구버전 병합 객체·신버전 runs 객체 혼재 가능)를 받아
    각 셀의 그리드 좌표를 복원한다. reservedUntilRow로 세로 병합이 점유 중인
    칸을 건너뛰는 알고리즘은 JsonToHwp.java의 addTableBlock과 동일하다(HTML
    표의 rowspan 렌더링과 같은 원리) — 두 방향(쓰기/읽기)이 같은 표현을
    쓰므로 재구성 로직도 같아야 한다."""
    n_rows = len(rows)
    norm_rows = []
    for row in rows:
        norm_row = []
        for cell in row:
            col_span = (cell.get("colSpan") or 1) if isinstance(cell, dict) else 1
            row_span = (cell.get("rowSpan") or 1) if isinstance(cell, dict) else 1
            norm_row.append({"runs": _cell_runs(cell), "colSpan": col_span, "rowSpan": row_span})
        norm_rows.append(norm_row)

    n_cols = sum(c["colSpan"] for c in norm_rows[0]) if norm_rows and norm_rows[0] else 0
    reserved_until_row = [-1] * n_cols
    placed = []
    for r in range(n_rows):
        col = 0
        cell_idx = 0
        row_cells = norm_rows[r]
        while col < n_cols and cell_idx < len(row_cells):
            if reserved_until_row[col] >= r:
                col += 1
                continue
            cell = row_cells[cell_idx]
            cell_idx += 1
            col_span = cell["colSpan"]
            row_span = cell["rowSpan"]
            placed.append({"runs": cell["runs"], "row": r, "col": col, "colSpan": col_span, "rowSpan": row_span})
            if row_span > 1:
                for cc in range(col, min(col + col_span, n_cols)):
                    reserved_until_row[cc] = r + row_span - 1
            col += col_span
    return placed, n_rows, n_cols


def blocks_to_docx(blocks: list[dict], out_path: Path) -> Path:
    from docx import Document

    doc = Document()
    for block in blocks:
        if block.get("type") == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            placed, n_rows, n_cols = _place_cells_with_spans(rows)
            if n_cols == 0:
                continue
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = "Table Grid"
            for item in placed:
                cell = table.cell(item["row"], item["col"])
                if item["colSpan"] > 1 or item["rowSpan"] > 1:
                    other = table.cell(
                        item["row"] + item["rowSpan"] - 1,
                        item["col"] + item["colSpan"] - 1,
                    )
                    cell = cell.merge(other)
                # cell.text = "..."(구버전) 대신 문단 블록과 같은 run 단위로
                # 직접 붙인다(DEC-051, 셀 안 문자 서식 보존). p.add_run(text)는
                # 임베드된 "\n"을 자동으로 <w:br/>로 바꿔주므로(python-docx
                # 자체 동작, 직접 확인) 여러 문단을 이어붙인 셀도 그대로 동작.
                p = cell.paragraphs[0]
                for run_dict in item["runs"]:
                    text = run_dict.get("text")
                    if not text:
                        continue
                    run = p.add_run(text)
                    _set_font(run)
                    _apply_run_style(run, run_dict)
        else:
            runs = [r for r in _runs_for(block) if r.get("text")]
            # 문단 전체가 공백뿐이면(실제 내용 없음) 문단 자체를 건너뛴다(기존 동작 유지).
            # 공백만 있는 개별 run은 단어 사이 구분자일 수 있어 그대로 둔다.
            if not any((r.get("text") or "").strip() for r in runs):
                continue
            p = doc.add_paragraph()
            for run_dict in runs:
                run = p.add_run(run_dict["text"])
                _set_font(run)
                _apply_run_style(run, run_dict)
            align = _align_map().get(block.get("align"))
            if align is not None:
                p.alignment = align
    doc.save(out_path)
    return out_path
