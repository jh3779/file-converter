"""PDF → DOCX 전용 변환기 — pdf.py의 공유 추출 프리미티브(_paragraph_candidates·
_iter_lines·_container_to_runs·_iter_visuals·_visual_to_dict)를 재사용해 DOCX
OOXML(w:framePr·w:pBdr·w:shd)을 직접 쓴다. pdf_pptx.py(PPTX 전용 슬라이드
작성 어댑터)와 대칭인 구조 — 구조 감사(2026-08)에서 세 관심사(공유 추출·
DOCX 작성·PPTX 작성)가 pdf.py 한 파일에 섞여 있던 것을 분리했다.
"""
from pathlib import Path

from .base import ConversionError
from .pdf import (
    _container_to_runs,
    _detect_alignment,
    _iter_lines,
    _iter_visuals,
    _paragraph_candidates,
    _underline_candidates,
    _visual_to_dict,
)


def pdf_to_docx(src: Path, tmpdir: Path) -> Path:
    """PDF → DOCX, 줄 단위로 원본과 같은 절대 위치에 배치해 재구성한다
    (DEC-037) — pdf_to_pptx(DEC-030, pdf_pptx.py)의 줄 단위 재구성과 같은
    원리를, DOCX의 레거시 "프레임" 기능(`w:framePr` — 문단을 페이지
    절대좌표에 고정, Word·LibreOffice 둘 다 지원)으로 구현한다. 이전에는
    문단을 순서대로 이어붙여 흐르게 하는 방식으로 단순화했었다(DEC-010) —
    이제 원본과 시각적으로 훨씬 가까운 배치를 낸다. HWP→DOCX는 아직 이
    방식으로 바뀌지 않았다(HWP 쪽 줄 단위 위치 정보를 새로 뽑아야 해서
    범위 밖 — hwp.py의 docx 변환은 여전히 blocks_to_docx 기반 흐르는 문서).

    pdf2docx(PyMuPDF)는 AGPL이라 사용 금지(DEC-007과 동일한 라이선스 원칙).

    **트레이드오프(사용자 확인 후 채택, 정직하게 문서화)**: 각 줄이 독립된
    프레임이 되므로, 결과 DOCX는 일반적인 "이어서 타이핑하면 자연스럽게
    다음 줄로 흐르는" 문서가 아니다 — 문장 중간에 텍스트를 추가해 그 줄
    프레임의 폭을 넘기면 다음 줄과 자연스럽게 안 이어진다(PDF→PPTX와
    동일한 트레이드오프 — 위치 정확도와 자유 편집 사이의 근본적 상충).

    이미지·벡터 도형(표 테두리 등)도 pdf_to_pptx(DEC-036)와 같은
    _iter_visuals/_visual_to_dict(pdf.py, 공유 헬퍼)로 뽑아 원래 위치·
    크기로 재구성한다(이미지·표 테두리 반영 개선) — python-docx엔 pptx의
    셰이프 API 같은 고수준 도형 삽입 기능이 없어(자유 배치 가능한 도형은
    python-docx가 지원 안 함), 이미 검증된 _set_frame_pr(줄 텍스트가
    쓰는 것과 동일)로 빈 문단을 이미지/도형 자리에 절대 위치시키고,
    이미지는 그 문단 안에 그림을 넣고, 사각형·직선은 문단 테두리(w:pBdr,
    DOCX가 지원하는 몇 안 되는 절대 배치 가능한 그리기 수단)로 표현한다.

    밑줄은 굵게/기울임과 달리 폰트 속성이 아니라 별도의 벡터 선(그림)으로
    그려지는 경우가 많아(PDF→DOCX 밑줄 감지 개선) 폰트 이름 휴리스틱과는
    별도로, 방금 뽑은 visuals 중 거의 수평인 선(_underline_candidates)이
    글자 바로 아래에 겹치는지 보고 판별한다(_char_is_underlined) — 이미
    밑줄로 쓰인 선은 run.underline=True(실제 DOCX 밑줄 서식)로 반영되고
    나면 표 테두리 도형으로 다시 그리면 겹쳐 보이므로 visuals 목록에서
    제외한다.

    **알려진 한계(정직하게 문서화)**: (1) 사각형이 아닌 곡선(LTCurve)은
    pdf_to_pptx처럼 다각형으로 정확히 근사하지 않고 bounding box 사각형
    테두리로 더 단순하게 근사한다 — python-docx가 pptx의 프리폼(자유 곡선)
    도형 API에 대응하는 기능이 없어 문단 테두리(사각형만 표현 가능)로
    타협했다(드문 경우, 표 테두리 대부분은 LTRect·LTLine). (2) 벡터·텍스트
    사이의 원래 z-순서는 완전히 보존되지 않는다(pdf_to_pptx와 동일한 이유 —
    도형을 먼저, 텍스트를 나중에 배치해 텍스트가 항상 위에 보이도록 단순화).
    (3) 디코딩 실패한 이미지는 조용히 건너뛴다(pdf_to_pptx와 동일).
    (4) 재구성에 항상 Noto Sans KR을 지정하므로(DEC-015와 같은 이유) 원본
    폰트와 글자 폭이 달라 드물게 줄바꿈이 살짝 밀릴 수 있다. (5) 기울임은
    한글 글꼴 대부분에 별도 이탤릭 글리프가 없어(CJK 타이포그래피 관행)
    감지되지 않는 경우가 흔함(pdf_to_pptx와 동일한 제약). (6) 밑줄 판정도
    위치 근접성만으로 추정하는 휴리스틱이라, 밑줄이 아니라 텍스트 바로
    아래 우연히 지나가는 다른 용도의 수평선(드문 경우)을 밑줄로 오판할
    가능성이 이론적으로 있다 — 반대로 표 테두리를 밑줄로 오판하면 그
    선이 표 테두리 도형에서는 빠지므로(사용됨 표시), 실사용 검증에서
    이런 오판이 나오면 겹침 판정 임계값(_UNDERLINE_MAX_GAP_PT 등)을
    조정할 것.
    """
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Emu, Pt

    from .docx_build import _align_map, _apply_run_style, _set_font

    EMU_PER_PT = 12700
    image_dir = tmpdir / "_pdf_docx_images"
    layout = _extract_pdf_line_layout(src, image_dir)
    if not layout:
        raise ConversionError("err.corrupted", "페이지 없음")

    doc = Document()

    def _size_section(section, page):
        section.page_width = Emu(round(page["width"] * EMU_PER_PT))
        section.page_height = Emu(round(page["height"] * EMU_PER_PT))
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, attr, Emu(0))

    _size_section(doc.sections[0], layout[0])
    current_size = (layout[0]["width"], layout[0]["height"])

    for page_index, page in enumerate(layout):
        if page_index > 0:
            page_size = (page["width"], page["height"])
            if page_size == current_size:
                doc.add_page_break()
            else:
                # 페이지 크기가 바뀌는 경우(스캔 첨부문서의 가로/세로 혼합 등)
                # — 단순 페이지 나눔만으로는 section의 page_width/height가
                # 첫 페이지 크기로 고정된 채라, 아래 vAnchor="page" 기준
                # y좌표(page_h - y1) 계산이 실제 렌더링 페이지 높이와 어긋나
                # 텍스트가 밀려 보이는 버그가 있었다(PR 콘텐츠 리뷰로 발견).
                # 크기가 실제로 바뀔 때만 새 섹션을 열어 그 페이지 실제
                # 크기로 맞춘다 — 흔한 동일 크기 다중 페이지는 기존과 동일하게
                # 단순 페이지 나눔을 그대로 쓴다.
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                _size_section(section, page)
                current_size = page_size
        page_h = page["height"]
        for visual in page["visuals"]:
            _add_visual_to_docx(doc, visual, page_h)
        for line in page["lines"]:
            x0, y0, x1, y1 = line["bbox"]
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            for run_dict in line["runs"]:
                run = p.add_run(run_dict["text"])
                _set_font(run)
                _apply_run_style(run, run_dict)
            align = _align_map().get(line.get("align"))
            if align is not None:
                p.alignment = align
            _set_frame_pr(p, x_pt=x0, y_pt=page_h - y1,
                          w_pt=max(x1 - x0, 1), h_pt=_text_frame_height_pt(y0, y1))

    out = tmpdir / (src.stem + ".docx")
    doc.save(out)
    return out


def _set_frame_pr(paragraph, x_pt: float, y_pt: float, w_pt: float, h_pt: float):
    """문단을 페이지 절대좌표에 고정한다(`w:framePr`) — DrawingML 플로팅
    도형보다 훨씬 단순한 레거시 기능이지만 Word·LibreOffice 둘 다 지원하며,
    이 용도(줄 단위 위치 재현)엔 충분하다. 스파이크로 실측: 지정한 x/y와
    LibreOffice로 렌더링한 PDF의 실제 텍스트 위치가 2pt 이내로 일치함을
    확인(hAnchor/vAnchor="page" 기준, 페이지 여백은 0으로 맞춰 좌표계를
    pdf_to_pptx의 EMU 변환과 동일하게 단순화). h_pt는 호출자가 이미
    필요한 여유를 계산해서 넘긴다고 가정한다(텍스트 줄은
    _text_frame_height_pt, 이미지·도형은 원래 크기 그대로 — 아래 참고)."""
    from docx.oxml.ns import qn

    TWIPS_PER_PT = 20
    pPr = paragraph._p.get_or_add_pPr()
    frame_pr = pPr.makeelement(qn("w:framePr"), {
        qn("w:w"): str(max(round(w_pt * TWIPS_PER_PT), 1)),
        qn("w:h"): str(max(round(h_pt * TWIPS_PER_PT), 1)),
        qn("w:hRule"): "exact",
        qn("w:hAnchor"): "page",
        qn("w:vAnchor"): "page",
        qn("w:x"): str(round(x_pt * TWIPS_PER_PT)),
        qn("w:y"): str(round(y_pt * TWIPS_PER_PT)),
        qn("w:wrap"): "none",
    })
    pPr.append(frame_pr)


# 번들 Noto Sans KR(Regular·Bold 둘 다 동일, fontTools로 직접 측정해 확인)의
# 실제 줄 높이 = (hhea/OS-2 usWin ascent + descent) / unitsPerEm = (1160+288)/1000
# = 1.448 — 즉 지정한 폰트 크기의 약 1.448배 높이가 있어야 글자(특히
# 내림선이 있는 g/p/디센더)가 위아래로 잘리지 않는다.
_NOTO_SANS_KR_LINE_HEIGHT_RATIO = 1.448


def _text_frame_height_pt(y0: float, y1: float) -> float:
    """텍스트 줄 프레임의 높이(pt)를 정한다 — pdf_to_docx 텍스트 클리핑
    결함 수정. PDF 원문 폰트가 무엇이었든 렌더링은 항상 Noto Sans KR로
    대체되는데(DEC-015), 이 폰트의 실제 줄 높이가 pdfminer가 준 원본
    bbox 높이(y1-y0, 원본 폰트 기준이라 대개 더 작음)보다 커서, 그
    bbox 높이 그대로 `w:framePr`(hRule="exact")을 걸면 초과분이 프레임
    밖으로 잘려 글자 대부분이 안 보이는 결함이 있었다(글꼴 크기·서식과
    무관하게 재현, DEC-055 검증 중 실제 LibreOffice 렌더링을 처음 육안
    확인해 발견). 대안으로 hRule="atLeast"(프레임이 필요하면 자동으로
    늘어남)도 시도했으나, 원본 PDF의 줄 간격이 원래 폰트 기준으로 촘촘한
    경우(흔함) 늘어난 프레임이 바로 다음 줄과 겹쳐 텍스트가 뒤섞여
    보이는 새 문제가 생겼다(직접 렌더링 비교로 확인) — 대신 필요한 높이를
    미리 계산해 hRule="exact"인 채로 정확히 그만큼만 키운다(오버랩·
    클리핑 둘 다 없음을 3줄짜리 촘촘한 간격으로 직접 확인). 이미지·도형
    (DEC-054)은 이 함수를 거치지 않고 원래 크기를 그대로 쓴다 — 폰트
    메트릭과 무관한 내용이라 키울 이유가 없다."""
    font_size = max(y1 - y0, 1)
    return font_size * _NOTO_SANS_KR_LINE_HEIGHT_RATIO


def _add_visual_to_docx(doc, visual: dict, page_h: float):
    """이미지·벡터 도형(사각형/직선/곡선) 하나를 원래 위치·크기의 빈
    문단으로 배치한다(pdf_to_docx, 이미지·표 테두리 반영 개선) — pptx의
    셰이프 API에 대응하는 기능이 python-docx엔 없어, 이미 검증된
    _set_frame_pr로 절대 위치를 잡고 내용물만 종류별로 다르게 채운다:
    이미지는 그림 run, 사각형·직선·곡선은 문단 테두리(w:pBdr)+채움
    (w:shd). 디코딩 실패 이미지는 pdf_to_pptx와 동일하게 조용히 건너뛴다
    (빈 위치 문단만 남기지 않도록 실패 시 문단 자체를 제거)."""
    from docx.shared import Pt

    x0, y0, x1, y1 = visual["bbox"]
    w_pt, h_pt = max(x1 - x0, 1), max(y1 - y0, 1)
    top_pt = page_h - y1

    if visual["kind"] == "image":
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        try:
            p.add_run().add_picture(visual["path"], width=Pt(w_pt), height=Pt(h_pt))
        except Exception:
            p._p.getparent().remove(p._p)
            return
        _set_frame_pr(p, x_pt=x0, y_pt=top_pt, w_pt=w_pt, h_pt=h_pt)
        return

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    linewidth_pt = max(visual.get("linewidth") or 0.75, 0.25)
    if visual["kind"] == "line":
        lx0, ly0 = visual["p0"]
        lx1, ly1 = visual["p1"]
        # 변 하나만 쓴다 — 위/아래(또는 좌/우) 두 변을 다 쓰면 얇은 프레임
        # 안에서 겹쳐 보여야 할 두 선이 육안으로 갈라져(이중선처럼) 보이는
        # 문제를 로컬에서 실제 LibreOffice 렌더링으로 확인해 수정.
        sides = ["right"] if abs(ly1 - ly0) > abs(lx1 - lx0) else ["bottom"]
    else:  # rect·curve(bounding box 사각형으로 근사)
        sides = ["top", "bottom", "left", "right"]
    if visual.get("stroke"):
        _set_paragraph_borders(p, sides, visual["stroke"], linewidth_pt)
    if visual.get("fill"):
        _set_paragraph_shading(p, visual["fill"])
    _set_frame_pr(p, x_pt=x0, y_pt=top_pt, w_pt=w_pt, h_pt=h_pt)


def _set_paragraph_borders(paragraph, sides: list[str], color_rgb: tuple[int, int, int], width_pt: float):
    """문단 테두리(`w:pBdr`)를 지정한 변에만 건다 — python-docx엔 표 셀
    테두리 API는 있어도 일반 문단 테두리 API가 없어 raw XML로 만든다.
    표 테두리(사각형)는 4변 전부, 직선은 방향에 맞는 변 1~2개만 받는다."""
    from docx.oxml.ns import qn

    EIGHTHS_PER_PT = 8
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    sz = str(max(round(width_pt * EIGHTHS_PER_PT), 2))
    color = "%02X%02X%02X" % color_rgb
    for side in sides:
        pBdr.append(pBdr.makeelement(qn(f"w:{side}"), {
            qn("w:val"): "single",
            qn("w:sz"): sz,
            qn("w:space"): "0",
            qn("w:color"): color,
        }))
    pPr.append(pBdr)


def _set_paragraph_shading(paragraph, color_rgb: tuple[int, int, int]):
    """문단 배경 채움(`w:shd`) — LTRect의 채움색(fill)을 표현한다."""
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "%02X%02X%02X" % color_rgb,
    }))


def _extract_pdf_line_layout(src: Path, image_dir: Path) -> list[dict]:
    """PDF → 페이지별 [{"width","height","lines":[{"bbox","runs","align"}],"visuals":[...]}] —
    pdf_to_docx 전용(DEC-037·이미지/표 테두리 반영 개선). pdf_to_pptx용
    _extract_pdf_layout()(pdf_pptx.py)과 이미지·도형(visuals) 추출은 같은
    헬퍼(pdf.py의 _iter_visuals/_visual_to_dict)를 그대로 재사용한다 —
    이제 DOCX도 이미지·표 테두리를 반영하므로 두 함수가 다시 갈라질
    이유는 "align"(문단 단위 정렬, PPTX 쪽은 안 씀)과 밑줄 감지(마찬가지로
    PPTX 쪽은 안 씀)뿐이라 유지한다. visuals는 줄을 훑기 전에 먼저 뽑는다
    — 밑줄 후보(_underline_candidates)를 문단 순회보다 앞서 계산해둬야
    각 줄의 _container_to_runs 호출에 넘겨줄 수 있기 때문이다. 정렬
    (DEC-040)은 컨테이너(원본 문단) 단위로 한 번만 판정해(_detect_alignment)
    그 문단에 속한 모든 줄에 같은 값을 붙인다 — 줄마다 독립된 프레임으로
    배치되지만(pdf_to_docx), 정렬 판정 자체는 여러 줄을 함께 봐야 하는
    문단 단위 신호이기 때문이다."""
    from pdfminer.high_level import extract_pages
    from pdfminer.image import ImageWriter
    from pdfminer.pdfdocument import PDFPasswordIncorrect
    from pdfminer.psparser import PSException

    image_dir.mkdir(parents=True, exist_ok=True)
    writer = ImageWriter(str(image_dir))

    pages = []
    try:
        for page in extract_pages(str(src)):
            visuals = [
                v for item in _iter_visuals(page)
                if (v := _visual_to_dict(item, writer, image_dir)) is not None
            ]
            underline_segments = _underline_candidates(visuals)
            lines = []
            for container in _paragraph_candidates(page):
                align = _detect_alignment(container, page.width)
                for line in _iter_lines(container):
                    runs = _container_to_runs(line, underline_segments=underline_segments)
                    if runs:
                        entry = {"bbox": line.bbox, "runs": runs}
                        if align:
                            entry["align"] = align
                        lines.append(entry)
            # 밑줄로 쓰인 선은 run.underline=True로 이미 반영됐으니
            # 표 테두리 도형(w:pBdr)으로 다시 그리면 겹쳐 보인다 — 제외.
            visuals = [v for v in visuals if not v.get("_used_as_underline")]
            pages.append({"width": page.width, "height": page.height, "lines": lines, "visuals": visuals})
    except PDFPasswordIncorrect:
        raise ConversionError("err.password")
    except (PSException, ValueError, OSError) as e:
        raise ConversionError("err.corrupted", str(e))
    return pages
